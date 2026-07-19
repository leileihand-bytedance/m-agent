from __future__ import annotations

from copy import deepcopy
import re
import shutil
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

from app.review.official_format_checker import load_official_format_rules
from skills.shenyinxie_news.schema import SelectedArticle
from skills.shenyinxie_news.selection import extract_markdown_front_matter


OUTPUT_FILENAME = "【深银协】微众银行{year}年{month}月第{monthly_issue}期信息动态.docx"
_REQUIRED_TEMPLATE_PLACEHOLDERS = (
    "{{TITLE}}",
    "{{PERIOD_RANGE}}",
    "{{ARTICLE_1_TITLE}}",
    "{{ARTICLE_1_BODY}}",
    "{{ARTICLE_1_SOURCE}}",
    "{{ARTICLE_2_TITLE}}",
    "{{ARTICLE_2_BODY}}",
    "{{ARTICLE_2_SOURCE}}",
    "{{ARTICLE_3_TITLE}}",
    "{{ARTICLE_3_BODY}}",
    "{{ARTICLE_3_SOURCE}}",
)
_WEB_PAGE_CHROME_PATTERNS = (
    re.compile(r"\d{4}-\d{1,2}-\d{1,2}(?:\s+\d{1,2}:\d{2}(?::\d{2})?)?"),
    re.compile(r"第\d{1,2}版[:：].*"),
    re.compile(r"人民日报\s+\d{4}年\d{1,2}月\d{1,2}日"),
)
_NEWSPAPER_METADATA_SUFFIX = re.compile(
    r"\s*《[^》]{1,30}》\s*[（(]\s*\d{4}年\d{1,2}月\d{1,2}日\s*"
    r"第\s*\d{1,2}\s*版\s*[）)]\s*$"
)
_WEB_PAGE_CHROME_EXACT = {
    "+",
    "-",
    "本版新闻",
    "Mon",
    "Tue",
    "Wed",
    "Thu",
    "Fri",
    "Sat",
    "Sun",
}
_RENDER_FONT_ALIASES = {
    "宋体": "Songti SC",
    "黑体": "Hiragino Sans GB",
}


def _default_template_path() -> Path:
    return Path(__file__).resolve().parent / "assets" / "shenyinxie-template.docx"


def write_shenyinxie_docx(
    *,
    title: str,
    period_start: date,
    period_end: date,
    issue_number: str,
    articles: list[SelectedArticle],
    output_dir: str | Path,
    template_path: Path | None = None,
) -> Path:
    """生成深银协动态 Word 文档。

    默认使用 Skill 自有的案例净化母版并按占位符替换。正式母版不可用时
    直接停止；只有调用方显式传入兼容模板时，才保留旧版代码生成兜底。
    """
    target_dir = Path(output_dir).resolve()
    target_dir.mkdir(parents=True, exist_ok=True)

    monthly_issue = _monthly_issue(period_start)
    target = target_dir / OUTPUT_FILENAME.format(
        year=period_start.year,
        month=period_start.month,
        monthly_issue=monthly_issue,
    )

    using_default_template = template_path is None
    template = template_path or _default_template_path()
    if template.exists() and _looks_like_real_template(template):
        _fill_template(template, target, title, period_start, period_end, issue_number, articles)
    elif using_default_template:
        raise RuntimeError("深银协正式模板不可用，已停止生成，避免交付错误版式")
    else:
        _build_from_scratch(target, title, period_start, period_end, issue_number, articles)

    return target


def _looks_like_real_template(template_path: Path) -> bool:
    """只有包含完整保真槽位的模板才允许进入原位替换流程。"""
    try:
        doc = Document(str(template_path))
        paragraph_texts = {paragraph.text.strip() for paragraph in doc.paragraphs}
        return all(placeholder in paragraph_texts for placeholder in _REQUIRED_TEMPLATE_PLACEHOLDERS)
    except Exception:
        return False


def _fill_template(
    template_path: Path,
    target: Path,
    title: str,
    period_start: date,
    period_end: date,
    issue_number: str,
    articles: list[SelectedArticle],
) -> None:
    """复制案例净化模板，在原位置填入标题、日期和逐段正文。"""
    shutil.copy(str(template_path), str(target))
    doc = Document(str(target))
    period_range = _format_period_range(period_start, period_end)

    _replace_paragraph_text(_find_placeholder(doc, "{{TITLE}}"), title)
    _replace_paragraph_text(
        _find_placeholder(doc, "{{PERIOD_RANGE}}"),
        f"（{period_range}）",
    )

    labels = ("一", "二", "三")
    for index in range(1, 4):
        heading = _find_placeholder(doc, f"{{{{ARTICLE_{index}_TITLE}}}}")
        body_marker = _find_placeholder(doc, f"{{{{ARTICLE_{index}_BODY}}}}")
        source = _find_placeholder(doc, f"{{{{ARTICLE_{index}_SOURCE}}}}")
        if index > len(articles):
            _remove_paragraph(heading)
            _remove_paragraph(body_marker)
            _remove_paragraph(source)
            continue

        article = articles[index - 1]
        _replace_paragraph_text(heading, f"【动态{labels[index - 1]}】{article.title}")
        _expand_body_paragraphs(
            body_marker,
            _article_body_paragraphs(article.body, article.title),
        )
        source_prototype = deepcopy(source._p)
        _set_source_paragraph(source, article)

        insertion_point = source
        if article.content_mode == "extract" and article.source_title:
            insertion_point = _insert_cloned_paragraph_after(
                insertion_point,
                source_prototype,
                f"原报道标题：{article.source_title}",
            )
        if article.content_mode == "extract" and article.editor_note:
            _insert_cloned_paragraph_after(
                insertion_point,
                source_prototype,
                article.editor_note,
            )

    doc.core_properties.title = title.strip() or "深银协动态"
    _apply_render_compatible_fonts(doc)
    doc.save(str(target))


def _build_from_scratch(
    target: Path,
    title: str,
    period_start: date,
    period_end: date,
    issue_number: str,
    articles: list[SelectedArticle],
) -> None:
    """无真实模板时，按公文格式新建文档。"""
    rules = load_official_format_rules()
    doc = Document()
    _apply_page_rules(doc, rules["page"])
    doc.core_properties.title = title.strip() or "深银协动态"

    _add_formatted_paragraph(doc, title, rules["roles"]["title"])
    period = _add_formatted_paragraph(
        doc,
        f"（{_format_period_range(period_start, period_end)}）",
        rules["roles"]["body"],
    )
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period.paragraph_format.first_line_indent = Pt(0)

    for idx, article in enumerate(articles, start=1):
        heading = f"【动态{['一', '二', '三'][idx - 1]}】{article.title}"
        _add_formatted_paragraph(doc, heading, rules["roles"]["heading1"])
        _add_article_block(doc, article, rules)

    _apply_render_compatible_fonts(doc)
    doc.save(str(target))


def _format_period_range(period_start: date, period_end: date) -> str:
    if period_start.year == period_end.year and period_start.month == period_end.month:
        return (
            f"{period_start.year}年{period_start.month}月{period_start.day}日-"
            f"{period_end.month}月{period_end.day}日"
        )
    if period_start.year == period_end.year:
        return (
            f"{period_start.year}年{period_start.month}月{period_start.day}日-"
            f"{period_end.month}月{period_end.day}日"
        )
    return (
        f"{period_start.year}年{period_start.month}月{period_start.day}日-"
        f"{period_end.year}年{period_end.month}月{period_end.day}日"
    )


def _monthly_issue(period_start: date) -> int:
    return 1 if period_start.day <= 15 else 2


def _find_placeholder(doc: Document, placeholder: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == placeholder:
            return paragraph
    raise ValueError(f"深银协模板缺少占位符：{placeholder}")


def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    for hyperlink in list(paragraph._p.xpath("./w:hyperlink")):
        hyperlink.getparent().remove(hyperlink)

    runs = list(paragraph.runs)
    keep = next((run for run in runs if run.text.strip()), runs[0] if runs else None)
    if keep is None:
        keep = paragraph.add_run()
    for run in runs:
        if run._element is not keep._element:
            run._element.getparent().remove(run._element)
    keep.text = text


def _remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def _insert_cloned_paragraph_after(
    paragraph: Paragraph,
    prototype: object,
    text: str,
) -> Paragraph:
    clone = deepcopy(prototype)
    paragraph._p.addnext(clone)
    inserted = Paragraph(clone, paragraph._parent)
    _replace_paragraph_text(inserted, text)
    return inserted


def _expand_body_paragraphs(marker: Paragraph, paragraphs: list[str]) -> None:
    paragraph_texts = paragraphs or [""]
    prototype = deepcopy(marker._p)
    _replace_paragraph_text(marker, paragraph_texts[0])
    current = marker
    for text in paragraph_texts[1:]:
        current = _insert_cloned_paragraph_after(current, prototype, text)


def _article_body_paragraphs(body: str, article_title: str = "") -> list[str]:
    _, body = extract_markdown_front_matter(body)
    paragraphs: list[str] = []
    for raw_line in body.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        clean = raw_line.replace("\u200b", "").replace("\ufeff", "").strip()
        clean = _NEWSPAPER_METADATA_SUFFIX.sub("", clean).strip()
        clean = re.sub(r"[\t ]+", " ", clean)
        if not clean or clean in _WEB_PAGE_CHROME_EXACT:
            continue
        if article_title and clean == article_title.strip():
            continue
        if any(pattern.fullmatch(clean) for pattern in _WEB_PAGE_CHROME_PATTERNS):
            continue
        paragraphs.append(clean)
    return paragraphs


def _apply_render_compatible_fonts(doc: Document) -> None:
    """只在生成副本中映射本机可用中文字体，避免渲染时正文消失。"""
    for root in (doc.styles.element, doc._element):
        for rfonts in root.xpath(".//w:rFonts"):
            for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
                key = qn(f"w:{attribute}")
                current = rfonts.get(key)
                replacement = _RENDER_FONT_ALIASES.get(current or "")
                if replacement:
                    rfonts.set(key, replacement)


def _format_publish_date(value: str) -> str:
    clean = value.strip()
    try:
        parsed = datetime.strptime(clean, "%Y-%m-%d").date()
    except ValueError:
        return clean
    return f"{parsed.year}年{parsed.month}月{parsed.day}日"


def _set_source_paragraph(paragraph: Paragraph, article: SelectedArticle) -> None:
    prefix = (
        f"来源：{article.media_name}　"
        f"发布时间：{_format_publish_date(article.publish_date)}　"
        "原文链接："
    )
    _replace_paragraph_text(paragraph, prefix)
    _append_hyperlink(paragraph, article.original_url)


def _append_hyperlink(paragraph: Paragraph, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend((color, underline))
    run.append(run_properties)

    text = OxmlElement("w:t")
    text.text = url
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_article_block(doc: Document, article: SelectedArticle, rules: dict[str, object]) -> None:
    body_rules = rules["roles"]["body"]

    # 正文
    for paragraph_text in _article_body_paragraphs(article.body, article.title):
        _add_formatted_paragraph(doc, paragraph_text, body_rules)

    # 摘编稿保留原报道标题，便于用户核对标题调整。
    if article.content_mode == "extract" and article.source_title:
        _add_formatted_paragraph(doc, f"原报道标题：{article.source_title}", body_rules)

    # 来源、日期、原文链接与摘编说明
    source = _add_formatted_paragraph(doc, "", body_rules)
    source.alignment = WD_ALIGN_PARAGRAPH.LEFT
    source.paragraph_format.first_line_indent = Pt(0)
    _set_source_paragraph(source, article)
    if article.content_mode == "extract" and article.editor_note:
        _add_formatted_paragraph(doc, article.editor_note, body_rules)


def _apply_page_rules(document: Document, rules: dict[str, object]) -> None:
    for section in document.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(float(rules["width_cm"]))
        section.page_height = Cm(float(rules["height_cm"]))
        section.top_margin = Cm(float(rules["top_margin_cm"]))
        section.bottom_margin = Cm(float(rules["bottom_margin_cm"]))
        section.left_margin = Cm(float(rules["left_margin_cm"]))
        section.right_margin = Cm(float(rules["right_margin_cm"]))


def _add_formatted_paragraph(document: Document, text: str, rules: dict[str, object]) -> object:
    paragraph = document.add_paragraph()
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
        if rules["alignment"] == "center"
        else WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    paragraph.paragraph_format.first_line_indent = Pt(float(rules["first_line_indent_pt"]))
    paragraph.paragraph_format.line_spacing = Pt(float(rules["line_spacing_pt"]))
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    run = paragraph.add_run(_clean_text(text))
    font_name = str(list(rules["fonts"])[0])
    run.font.name = font_name
    run.font.size = Pt(float(rules["size_pt"]))
    run.font.bold = bool(rules["bold"])
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    return paragraph


def _clean_text(text: str) -> str:
    clean = text.strip()
    clean = re.sub(r"^#{1,6}\s*", "", clean)
    clean = clean.replace("**", "").replace("__", "")
    return clean.strip()


__all__ = ["write_shenyinxie_docx"]
