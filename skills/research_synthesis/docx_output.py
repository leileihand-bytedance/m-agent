from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.review.official_format_checker import load_official_format_rules


INTRO_NOTE = "【备注：请根据实际报送对象和通知要求补充报告开头。】"
ENDING_NOTE = "【备注：请根据实际报送要求补充报告结尾、附件、联系人、落款和日期。】"
OUTPUT_FILENAME = "综合调研材料初稿.docx"
_H1_RE = re.compile(r"^[一二三四五六七八九十百零〇]+、")
_H2_RE = re.compile(r"^（[一二三四五六七八九十百零〇]+）")
_H3_RE = re.compile(r"^\d+[.．、](?!\d)")


def write_research_synthesis_docx(*, title: str, body: str, output_dir: str | Path) -> Path:
    target_dir = Path(output_dir).resolve()
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / OUTPUT_FILENAME
    rules = load_official_format_rules()
    document = Document()
    _apply_page_rules(document, rules["page"])
    document.core_properties.title = title.strip() or "综合调研材料"

    clean_title = _clean_line(title) or "综合调研材料"
    _add_formatted_paragraph(document, clean_title, rules["roles"]["title"])
    _add_formatted_paragraph(document, INTRO_NOTE, rules["roles"]["body"])
    for line in _body_lines(body, title=clean_title):
        role = _role_for_text(line)
        _add_formatted_paragraph(document, line, rules["roles"][role])
    _add_formatted_paragraph(document, ENDING_NOTE, rules["roles"]["body"])
    document.save(target)
    return target


def _apply_page_rules(document: Document, rules: dict[str, object]) -> None:
    for section in document.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(float(rules["width_cm"]))
        section.page_height = Cm(float(rules["height_cm"]))
        section.top_margin = Cm(float(rules["top_margin_cm"]))
        section.bottom_margin = Cm(float(rules["bottom_margin_cm"]))
        section.left_margin = Cm(float(rules["left_margin_cm"]))
        section.right_margin = Cm(float(rules["right_margin_cm"]))


def _body_lines(body: str, *, title: str) -> list[str]:
    lines = [_clean_line(line) for line in body.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    clean_lines = [line for line in lines if line]
    if clean_lines and clean_lines[0] == title:
        clean_lines = clean_lines[1:]
    return [line for line in clean_lines if line not in {INTRO_NOTE, ENDING_NOTE}]


def _clean_line(line: str) -> str:
    clean = line.strip()
    clean = re.sub(r"^#{1,6}\s*", "", clean)
    clean = clean.replace("**", "").replace("__", "")
    return clean.strip()


def _role_for_text(text: str) -> str:
    if _H1_RE.match(text):
        return "heading1"
    if _H2_RE.match(text):
        return "heading2"
    if _H3_RE.match(text):
        return "heading3"
    return "body"


def _add_formatted_paragraph(document: Document, text: str, rules: dict[str, object]) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
        if rules["alignment"] == "center"
        else WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    paragraph.paragraph_format.first_line_indent = Pt(float(rules["first_line_indent_pt"]))
    paragraph.paragraph_format.line_spacing = Pt(float(rules["line_spacing_pt"]))
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    run = paragraph.add_run(text)
    font_name = str(list(rules["fonts"])[0])
    run.font.name = font_name
    run.font.size = Pt(float(rules["size_pt"]))
    run.font.bold = bool(rules["bold"])
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


__all__ = [
    "ENDING_NOTE",
    "INTRO_NOTE",
    "OUTPUT_FILENAME",
    "write_research_synthesis_docx",
]
