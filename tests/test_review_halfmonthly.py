"""半月报审核模块测试."""

from __future__ import annotations

import asyncio
from datetime import date
from pathlib import Path

import pytest

from app.review.document_type import detect_document_type, DocumentType, document_type_label
from app.review.halfmonthly_reviewer import (
    parse_halfmonthly_date_range,
    _check_date_range,
    _check_leader_title,
    _check_section_order,
    review_halfmonthly,
    HALFMONTHLY_SECTION_TITLES,
)
from app.review.output_formatter import format_review_result
from app.review.reviewer import ReviewResult, Finding
from app.review.error_marker import mark_errors_in_docx


def _make_minimal_docx(path: Path, paragraphs: list[str]) -> None:
    """构造一个最小 .docx."""
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))


def _marked_texts(path: Path) -> list[str]:
    """读取 docx 中所有红色 run 的文本."""
    from docx import Document
    from docx.shared import RGBColor

    doc = Document(str(path))
    result: list[str] = []
    for paragraph in doc.paragraphs:
        parts: list[str] = []
        for run in paragraph.runs:
            if run.font.color and run.font.color.rgb == RGBColor(255, 0, 0):
                parts.append(run.text)
        if parts:
            result.append("".join(parts))
    return result


# ============================================================
# 1. 文档类型识别
# ============================================================

def test_detect_document_type_by_filename():
    assert detect_document_type("信息动态半月报26年4月第1期.docx", []) == DocumentType.HALF_MONTHLY
    assert detect_document_type("微众银行信息内参周报2026年第26期.docx", []) == DocumentType.NEI_CAN
    assert detect_document_type("某周报.docx", []) == DocumentType.NEI_CAN
    assert detect_document_type(None, ["内部资料", "微众银行信息动态半月报"]) == DocumentType.HALF_MONTHLY
    assert detect_document_type(None, ["内部资料", "微众银行信息内参周报"]) == DocumentType.NEI_CAN


def test_document_type_label():
    assert document_type_label(DocumentType.HALF_MONTHLY) == "半月报"
    assert document_type_label(DocumentType.NEI_CAN) == "内参周报"


# ============================================================
# 2. 时间范围解析
# ============================================================

def test_parse_halfmonthly_date_range_with_parens():
    paragraphs = [
        "内部资料",
        "微众银行信息动态半月报",
        "（2026年4月1日-4月15日）",
        "业务动态及成果",
    ]
    rng = parse_halfmonthly_date_range(paragraphs)
    assert rng is not None
    assert rng.start == date(2026, 4, 1)
    assert rng.end == date(2026, 4, 15)


def test_parse_halfmonthly_date_range_without_parens():
    paragraphs = [
        "内部资料",
        "信息动态半月报",
        "2026年4月1日-4月15日",
    ]
    rng = parse_halfmonthly_date_range(paragraphs)
    assert rng is not None
    assert rng.start == date(2026, 4, 1)
    assert rng.end == date(2026, 4, 15)


def test_parse_halfmonthly_date_range_not_found():
    paragraphs = ["内部资料", "微众银行信息内参周报"]
    assert parse_halfmonthly_date_range(paragraphs) is None


# ============================================================
# 3. 日期范围代码预检
# ============================================================

def test_check_date_range_out_of_range():
    from app.review.halfmonthly_reviewer import DateRange
    paragraphs = [
        "内部资料",
        "半月报",
        "2026年4月1日-4月15日",
        "业务动态及成果",
        "3月28日,我行发生某业务。",
        "4月20日,我行发生另一业务。",
        "4月3日,我行正常业务。",
    ]
    rng = DateRange(start=date(2026, 4, 1), end=date(2026, 4, 15))
    findings = _check_date_range(paragraphs, rng)
    indices = {f.paragraph_index for f in findings}
    assert 4 in indices
    assert 5 in indices
    assert 6 not in indices


def test_check_date_range_skips_section_titles():
    from app.review.halfmonthly_reviewer import DateRange
    paragraphs = [
        "业务动态及成果",
        "4月3日,正常业务。",
    ]
    rng = DateRange(start=date(2026, 4, 1), end=date(2026, 4, 15))
    findings = _check_date_range(paragraphs, rng)
    assert all(f.paragraph_index != 0 for f in findings)


def test_check_section_order_allows_in_order_subset():
    paragraphs = [
        "内部资料",
        "微众银行信息动态半月报",
        "（2026年4月1日-4月15日）",
        "工作动态及成果",
        "4月3日,某项目推进完成阶段成果。",
        "获得资质与荣誉",
        "4月8日,我行获得某项行业奖项。",
    ]
    findings = _check_section_order(paragraphs)
    assert findings == []


def test_check_section_order_reports_out_of_order_section_title():
    paragraphs = [
        "内部资料",
        "微众银行信息动态半月报",
        "（2026年4月1日-4月15日）",
        "行外联络及交流",
        "4月9日,外部机构来访交流。",
        "获得资质与荣誉",
        "4月10日,我行获得某项资质。",
    ]
    findings = _check_section_order(paragraphs)

    assert len(findings) == 1
    assert findings[0].rule_id == "halfmonthly-section-order"
    assert findings[0].paragraph_index == 5
    assert findings[0].target_text == "获得资质与荣誉"
    assert "应排在'行外联络及交流'之前" in findings[0].description


# ============================================================
# 4. 半月报审核(mock LLM)
# ============================================================

class _FakeBlock:
    def __init__(self, text: str):
        self.text = text


class _FakeMessage:
    def __init__(self, text: str):
        self.content = [_FakeBlock(text)]


class _FakeMessages:
    def __init__(self, counter: dict[str, int], text: str):
        self._counter = counter
        self._text = text

    def create(self, **_: object) -> _FakeMessage:
        self._counter["calls"] += 1
        return _FakeMessage(self._text)


class _FakeClient:
    def __init__(self, counter: dict[str, int], text: str):
        self.messages = _FakeMessages(counter, text)


def test_review_halfmonthly_mock_llm(monkeypatch):
    counter = {"calls": 0}
    fake_output = '{"issues": []}'
    monkeypatch.setattr(
        "app.review.halfmonthly_reviewer.build_anthropic_client",
        lambda: (_FakeClient(counter, fake_output), "fake-model"),
    )

    paragraphs = [
        "内部资料",
        "微众银行信息动态半月报",
        "（2026年4月1日-4月15日）",
        "业务动态及成果",
        "4月3日,我行顺利落地某业务。",
    ]
    result = asyncio.run(review_halfmonthly(paragraphs, "", "半月报.docx"))

    # LLM 返回 0 条时会再试一次,所以最多 2 次
    assert counter["calls"] <= 2
    assert counter["calls"] >= 1
    assert not any(f.rule_id.startswith("__") for f in result.findings)
    assert result.filename == "半月报.docx"


def test_review_halfmonthly_date_mismatch_found_without_llm(monkeypatch):
    """日期超出范围由代码预检发现,即使 LLM 返回空也能命中."""
    counter = {"calls": 0}
    fake_output = '{"issues": []}'
    monkeypatch.setattr(
        "app.review.halfmonthly_reviewer.build_anthropic_client",
        lambda: (_FakeClient(counter, fake_output), "fake-model"),
    )

    paragraphs = [
        "内部资料",
        "微众银行信息动态半月报",
        "（2026年4月1日-4月15日）",
        "业务动态及成果",
        "3月28日,我行发生某业务。",
    ]
    result = asyncio.run(review_halfmonthly(paragraphs, "", "半月报.docx"))

    date_mismatches = [f for f in result.findings if f.rule_id == "halfmonthly-date-mismatch"]
    assert len(date_mismatches) == 1
    assert date_mismatches[0].paragraph_index == 4


def test_review_halfmonthly_section_order_found_without_llm(monkeypatch):
    """一级标题乱序由代码预检发现,即使 LLM 返回空也能命中."""
    counter = {"calls": 0}
    fake_output = '{"issues": []}'
    monkeypatch.setattr(
        "app.review.halfmonthly_reviewer.build_anthropic_client",
        lambda: (_FakeClient(counter, fake_output), "fake-model"),
    )

    paragraphs = [
        "内部资料",
        "微众银行信息动态半月报",
        "（2026年4月1日-4月15日）",
        "行外联络及交流",
        "4月9日,外部机构来访交流。",
        "获得资质与荣誉",
        "4月10日,我行获得某项资质。",
    ]
    result = asyncio.run(review_halfmonthly(paragraphs, "", "半月报.docx"))

    section_order_findings = [f for f in result.findings if f.rule_id == "halfmonthly-section-order"]
    assert len(section_order_findings) == 1
    assert section_order_findings[0].paragraph_index == 5


def test_review_halfmonthly_huang_missing_party_title_found_without_llm(monkeypatch):
    """明确党内职务口径下,黄黎明漏写党委副书记应直接命中."""
    counter = {"calls": 0}
    fake_output = '{"issues": []}'
    monkeypatch.setattr(
        "app.review.halfmonthly_reviewer.build_anthropic_client",
        lambda: (_FakeClient(counter, fake_output), "fake-model"),
    )

    paragraphs = [
        "内部资料",
        "微众银行信息动态半月报",
        "（2026年4月1日-4月15日）",
        "行内重要会议",
        "4月3日,李南青,党委书记，万军,党委委员、监事会主席，黄黎明,行长出席某活动。",
    ]
    result = asyncio.run(review_halfmonthly(paragraphs, "", "半月报.docx"))

    leader_findings = [f for f in result.findings if f.rule_id == "halfmonthly-leader-title"]
    assert len(leader_findings) == 1
    assert "黄黎明" in leader_findings[0].description
    assert "党委副书记" in leader_findings[0].description


# ============================================================
# 5. 格式化输出
# ============================================================

def test_format_review_result_shows_halfmonthly_label():
    result = ReviewResult(
        findings=[],
        total_rules=10,
        passed_rules=10,
        filename="半月报.docx",
    )
    output = format_review_result(result, "半月报.docx", doc_type=DocumentType.HALF_MONTHLY)
    assert "(半月报)" in output


# ============================================================
# 6. 领导职务与排序规范
# ============================================================

def test_check_leader_title_suppressed_titles():
    """检测默认不应出现的职务片段."""
    paragraphs = [
        "内部资料",
        "4月3日,李南青,党委书记、首席合规官出席某会议。",
        "4月5日,黄黎明,行长、党委副书记接待来访。",
    ]
    findings = _check_leader_title(paragraphs)
    rule_ids = [f.rule_id for f in findings]
    assert rule_ids.count("halfmonthly-leader-title") == 2
    descriptions = " ".join(f.description for f in findings)
    assert "首席合规官" in descriptions
    assert "党委副书记" in descriptions


def test_check_leader_title_allows_huang_party_title_in_party_title_style():
    """整体采用党内职务口径时,黄黎明写党委副书记不报错."""
    paragraphs = [
        "内部资料",
        "4月3日,李南青,党委书记，黄黎明,行长、党委副书记，万军,党委委员、监事会主席出席会议。",
    ]
    findings = _check_leader_title(paragraphs)
    huang_findings = [f for f in findings if "黄黎明" in f.original_text and "党委副书记" in f.description]
    assert huang_findings == []


def test_check_leader_title_huang_missing_party_title_is_allowed_in_default_style():
    """默认口径下,黄黎明不写党委副书记不报错."""
    paragraphs = [
        "内部资料",
        "4月3日,黄黎明,行长和陈峭,副行长出席某活动。",
    ]
    findings = _check_leader_title(paragraphs)
    huang_missing = [f for f in findings if "黄黎明" in f.description and "补充" in f.description]
    assert huang_missing == []


def test_check_leader_title_internal_party_title_triggers_other_reminders():
    """除李南青外,已有人写党内职务时,其他相关领导要补齐."""
    paragraphs = [
        "内部资料",
        "4月3日,李南青,党委书记，黄黎明,行长、党委副书记，方震宇,副行长，王立鹏,行长助理出席某活动。",
    ]
    findings = _check_leader_title(paragraphs)
    other_party_reminders = [
        f for f in findings
        if "方震宇" in f.description or "王立鹏" in f.description or "公立" in f.description or "江旻" in f.description
    ]
    assert len(other_party_reminders) == 2
    assert any("方震宇" in f.description and "党委委员" in f.description for f in other_party_reminders)
    assert any("王立鹏" in f.description and "党委委员" in f.description for f in other_party_reminders)


def test_check_leader_title_li_alone_does_not_trigger_other_reminders():
    """只有李南青写党委书记时,不触发其他人的党内职务补齐提醒."""
    paragraphs = [
        "内部资料",
        "4月3日,李南青,党委书记，黄黎明,行长，方震宇,副行长出席某活动。",
    ]
    findings = _check_leader_title(paragraphs)
    missing_party = [f for f in findings if "补充" in f.description]
    assert missing_party == []


def test_check_leader_title_requires_huang_party_title_in_party_title_style():
    """整体采用党内职务口径时,黄黎明未写党委副书记应报错."""
    paragraphs = [
        "内部资料",
        "4月3日,李南青,党委书记，万军,党委委员、监事会主席，黄黎明,行长出席某活动。",
    ]
    findings = _check_leader_title(paragraphs)
    missing_party = [f for f in findings if "黄黎明" in f.description and "补充" in f.description]

    assert len(missing_party) == 1
    assert "党委副书记" in missing_party[0].description


def test_check_leader_title_external_party_title_triggers_internal_reminders():
    """第三方人员写党内职务时,也会触发行内领导补齐提醒."""
    paragraphs = [
        "内部资料",
        "4月3日,黄黎明,行长，方震宇,副行长会见某集团党委书记张三一行。",
    ]
    findings = _check_leader_title(paragraphs)
    missing_party = [f for f in findings if "补充" in f.description]

    assert len(missing_party) == 2
    assert any("黄黎明" in f.description and "党委副书记" in f.description for f in missing_party)
    assert any("方震宇" in f.description and "党委委员" in f.description for f in missing_party)


def test_check_leader_title_allows_huang_party_title_when_external_party_title_present():
    """第三方已采用党内职务口径时,黄黎明写党委副书记不报错."""
    paragraphs = [
        "内部资料",
        "4月3日,黄黎明,行长、党委副书记会见某集团党委书记张三一行。",
    ]
    findings = _check_leader_title(paragraphs)
    huang_findings = [f for f in findings if "黄黎明" in f.original_text and "党委副书记" in f.description]
    assert huang_findings == []


def test_check_leader_title_order_error():
    """检测同一段落领导排序错误."""
    paragraphs = [
        "内部资料",
        "4月3日,方震宇和陈峭共同出席某活动。",
    ]
    findings = _check_leader_title(paragraphs)
    order_errors = [f for f in findings if "排序错误" in f.description]
    assert len(order_errors) == 1
    assert "陈峭不应在方震宇之前" in order_errors[0].description


def test_check_leader_title_order_correct():
    """正确排序不报."""
    paragraphs = [
        "内部资料",
        "4月3日,陈峭和方震宇共同出席某活动。",
    ]
    findings = _check_leader_title(paragraphs)
    order_errors = [f for f in findings if "排序错误" in f.description]
    assert len(order_errors) == 0


def test_leader_title_label_in_formatter():
    """formatter 正确显示领导职务规范标签."""
    result = ReviewResult(
        findings=[
            Finding(
                rule_id="halfmonthly-leader-title",
                paragraph_index=1,
                line_number=2,
                original_text="李南青,党委书记、首席合规官",
                description="李南青默认不应出现'首席合规官'",
            ),
        ],
        total_rules=11,
        passed_rules=10,
        filename="半月报.docx",
    )
    output = format_review_result(result, "半月报.docx", doc_type=DocumentType.HALF_MONTHLY)
    assert "半月报领导职务规范" in output


def test_section_order_label_in_formatter():
    """formatter 正确显示半月报板块顺序标签."""
    result = ReviewResult(
        findings=[
            Finding(
                rule_id="halfmonthly-section-order",
                paragraph_index=5,
                line_number=6,
                original_text="获得资质与荣誉",
                description="一级标题顺序错误：'获得资质与荣誉'应排在'行外联络及交流'之前",
                target_text="获得资质与荣誉",
            ),
        ],
        total_rules=12,
        passed_rules=11,
        filename="半月报.docx",
    )
    output = format_review_result(result, "半月报.docx", doc_type=DocumentType.HALF_MONTHLY)
    assert "半月报板块顺序" in output

# ============================================================
# 7. 标红定位能力
# ============================================================

def test_check_date_range_finding_has_target_text():
    from app.review.halfmonthly_reviewer import DateRange
    paragraphs = [
        "内部资料",
        "半月报",
        "2026年4月1日-4月15日",
        "业务动态及成果",
        "3月28日,我行发生某业务。",
    ]
    rng = DateRange(start=date(2026, 4, 1), end=date(2026, 4, 15))
    findings = _check_date_range(paragraphs, rng)

    assert len(findings) == 1
    assert findings[0].target_text == "3月28日"


def test_check_leader_title_suppressed_finding_has_target_text():
    paragraphs = [
        "内部资料",
        "4月3日,李南青,党委书记、首席合规官出席某会议。",
    ]
    findings = _check_leader_title(paragraphs)
    suppressed = [f for f in findings if "首席合规官" in f.description]

    assert len(suppressed) == 1
    assert suppressed[0].target_text == "首席合规官"


def test_check_leader_title_order_finding_has_target_text():
    paragraphs = [
        "内部资料",
        "4月3日,方震宇和陈峭共同出席某活动。",
    ]
    findings = _check_leader_title(paragraphs)
    order_errors = [f for f in findings if "排序错误" in f.description]

    assert len(order_errors) == 1
    assert order_errors[0].target_text == "陈峭"


def test_build_halfmonthly_prompt_mentions_target_text():
    from app.review.halfmonthly_reviewer import _build_halfmonthly_prompt

    prompt = _build_halfmonthly_prompt(
        "rules",
        ["业务动态及成果", "4月3日,某业务落地。"],
        "半月报.docx",
        None,
    )
    assert '"target_text"' in prompt
    assert "精确定位标红位置" in prompt


def test_mark_errors_in_docx_highlights_halfmonthly_finding(tmp_path: Path):
    """半月报 finding 带 target_text 时,应能在 docx 中精确标红."""
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _make_minimal_docx(input_path, ["4月3日,李南青,党委书记、首席合规官出席某会议。"])

    findings = [
        Finding(
            rule_id="halfmonthly-leader-title",
            paragraph_index=0,
            line_number=1,
            original_text="4月3日,李南青,党委书记、首席合规官出席某会议。",
            description="李南青默认不应出现'首席合规官'",
            target_text="首席合规官",
        ),
    ]
    mark_errors_in_docx(input_path, output_path, findings)

    marked = _marked_texts(output_path)
    assert any("首席合规官" in text for text in marked)
