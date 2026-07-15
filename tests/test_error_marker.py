"""通用审核错误标记器测试."""

from __future__ import annotations

import zipfile
from pathlib import Path

from app.review.error_marker import (
    mark_errors_in_docx,
    _get_search_key,
    _build_comment_text,
)
from app.review.format_checker import check_quote_pair
from app.review.parser import parse_docx
from app.review.reviewer import Finding


def _make_minimal_docx(path: Path, paragraphs: list[str]) -> None:
    """构造一个最小 .docx."""
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))


def _inject_null_relationship(path: Path) -> None:
    """向 .docx 的关系文件中插入一个 NULL 关系."""
    tmp_path = path.with_name(f"{path.stem}_patched{path.suffix}")

    with zipfile.ZipFile(path, "r") as zin:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "_rels/.rels":
                    rels = data.decode("utf-8")
                    rels = rels.replace(
                        "</Relationships>",
                        '<Relationship Id="rIdNULL" '
                        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/null" '
                        'Target="NULL" /></Relationships>',
                    )
                    data = rels.encode("utf-8")
                zout.writestr(info, data)

    tmp_path.replace(path)


def _marked_texts(path: Path) -> list[str]:
    """读取 docx 中所有红色 run 的文本."""
    from docx import Document
    from docx.shared import RGBColor

    doc = Document(str(path))
    result: list[str] = []
    for paragraph in doc.paragraphs:
        parts: list[str] = []
        for run in paragraph.runs:
            if run.font.color and run.font.color.rgb == RGBColor(255, 0, 0):
                parts.append(run.text)
        if parts:
            result.append("".join(parts))
    return result


def _marked_offsets(path: Path) -> list[tuple[int, int, str]]:
    """读取红色 run 的段号、字符起点和文本."""
    from docx import Document
    from docx.shared import RGBColor

    result: list[tuple[int, int, str]] = []
    for paragraph_index, paragraph in enumerate(Document(str(path)).paragraphs):
        offset = 0
        for run in paragraph.runs:
            if run.font.color and run.font.color.rgb == RGBColor(255, 0, 0):
                result.append((paragraph_index, offset, run.text))
            offset += len(run.text)
    return result


def test_get_search_key_uses_target_text():
    f = Finding(
        rule_id="general-typo",
        paragraph_index=0,
        line_number=1,
        original_text="本周布署了工作。",
        description="'部署'误写为'布署'",
        target_text="布署",
    )
    assert _get_search_key(f) == "布署"


def test_get_search_key_falls_back_to_description():
    f = Finding(
        rule_id="general-typo",
        paragraph_index=0,
        line_number=1,
        original_text="本周布署了工作。",
        description="'部署'误写为'布署'",
    )
    assert _get_search_key(f) == "布署"


def test_mark_errors_highlights_exact_target_for_typo(tmp_path: Path):
    """错别字只标红错误词，避免长句大面积标红."""
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _make_minimal_docx(input_path, ["本周布署了工作。下周继续推进。"])

    findings = [
        Finding(
            rule_id="general-typo",
            paragraph_index=0,
            line_number=1,
            original_text="本周布署了工作。",
            description="'部署'误写为'布署'",
            target_text="布署",
        ),
    ]

    mark_errors_in_docx(input_path, output_path, findings)
    assert output_path.exists()

    marked = _marked_texts(output_path)
    assert marked == ["布署"]


def test_mark_errors_highlights_exact_target_for_multi_file_issue(tmp_path: Path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _make_minimal_docx(input_path, ["请填写附件1《议案意见反馈表》。"])
    findings = [
        Finding(
            rule_id="multi-file-attachment-name-mismatch",
            paragraph_index=0,
            line_number=1,
            original_text="请填写附件1《议案意见反馈表》。",
            description="实际上传的附件1标题不一致",
            target_text="附件1",
        )
    ]

    mark_errors_in_docx(input_path, output_path, findings)

    assert _marked_texts(output_path) == ["附件1"]


def test_mark_errors_highlights_exact_target_for_punctuation(tmp_path: Path):
    """标点错误只标红错误标点."""
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _make_minimal_docx(input_path, ["会议讨论了A,B,C三个议题。"])

    findings = [
        Finding(
            rule_id="general-punctuation",
            paragraph_index=0,
            line_number=1,
            original_text="会议讨论了A,B,C三个议题。",
            description="中文句子里出现英文逗号",
            target_text=",",
        ),
    ]

    mark_errors_in_docx(input_path, output_path, findings)
    assert output_path.exists()

    marked = _marked_texts(output_path)
    assert marked == [","]


def test_mark_errors_preserves_significant_space_after_punctuation(tmp_path: Path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    paragraph = "请中原银行、 四川银行作答。"
    _make_minimal_docx(input_path, [paragraph])

    findings = [
        Finding(
            rule_id="general-punctuation",
            paragraph_index=0,
            line_number=1,
            original_text=paragraph,
            description="顿号后有多余空格，应删除该空格",
            target_text="、 ",
        ),
    ]

    mark_errors_in_docx(input_path, output_path, findings)

    assert _get_search_key(findings[0]) == "、 "
    assert _marked_texts(output_path) == ["、 "]


def test_mark_errors_consecutive_punctuation(tmp_path: Path):
    """连续标点应标红错误标点及前面半句话."""
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _make_minimal_docx(input_path, ["今天开会了。。明天继续。"])

    findings = [
        Finding(
            rule_id="general-punctuation",
            paragraph_index=0,
            line_number=1,
            original_text="今天开会了。。明天继续。",
            description="连续相同标点：'。。'",
            target_text="。。",
        ),
    ]

    mark_errors_in_docx(input_path, output_path, findings)
    assert output_path.exists()

    marked = _marked_texts(output_path)
    assert marked == ["。。"]


def test_mark_errors_splits_runs_for_realistic_paragraph(tmp_path: Path):
    """真实文档中一个段落常由多个 run 组成,应正确拆分并标红."""
    from docx import Document

    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    paragraph = doc.add_paragraph()
    for fragment in ["本周", "重点", "布署", "了工作。", "下周继续"]:
        run = paragraph.add_run(fragment)
        run.bold = True
    doc.save(str(input_path))

    findings = [
        Finding(
            rule_id="general-typo",
            paragraph_index=0,
            line_number=1,
            original_text="本周重点布署了工作。",
            description="'部署'误写为'布署'",
            target_text="布署",
        ),
    ]

    mark_errors_in_docx(input_path, output_path, findings)
    marked = _marked_texts(output_path)
    assert marked == ["布署"]


def test_mark_errors_skips_invalid_paragraph_index(tmp_path: Path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _make_minimal_docx(input_path, ["只有一段。"])

    findings = [
        Finding(
            rule_id="general-typo",
            paragraph_index=99,
            line_number=100,
            original_text="越界",
            description="越界测试",
            target_text="越界",
        ),
    ]

    mark_errors_in_docx(input_path, output_path, findings)
    assert output_path.exists()
    assert _marked_texts(output_path) == []


def test_inject_null_relationship_replaces_rels_entry_without_duplicates(tmp_path: Path):
    input_path = tmp_path / "null_rel.docx"
    _make_minimal_docx(input_path, ["测试段落。"])

    _inject_null_relationship(input_path)

    with zipfile.ZipFile(input_path) as zf:
        assert zf.namelist().count("_rels/.rels") == 1


def test_parse_docx_handles_null_relationship(tmp_path: Path):
    """解析包含指向 NULL 的 OPC 关系的 .docx 时不抛异常."""
    from docx import Document

    input_path = tmp_path / "null_rel.docx"
    doc = Document()
    doc.add_paragraph("测试段落。")
    doc.save(str(input_path))
    _inject_null_relationship(input_path)

    parsed = parse_docx(input_path)
    assert parsed.paragraphs == ["测试段落。"]


def test_parse_docx_large_table_is_stable_and_keeps_every_cell(tmp_path: Path):
    """大表格不能因单元格对象 ID 复用而随机漏段."""
    from docx import Document

    input_path = tmp_path / "large_table.docx"
    doc = Document()
    table = doc.add_table(rows=30, cols=4)
    expected = []
    for row_idx in range(30):
        for col_idx in range(4):
            text = f"cell-{row_idx}-{col_idx}"
            table.cell(row_idx, col_idx).text = text
            expected.append(text)
    doc.save(str(input_path))

    parsed_runs = [parse_docx(input_path).paragraphs for _ in range(3)]

    assert parsed_runs == [expected, expected, expected]


def test_mark_errors_handles_null_relationship_docx(tmp_path: Path):
    """对包含 NULL 关系的 .docx 也能成功生成标注文档."""
    from docx import Document

    input_path = tmp_path / "null_rel.docx"
    output_path = tmp_path / "null_rel_marked.docx"

    doc = Document()
    doc.add_paragraph("中国人民银深圳市中心支行：")
    doc.save(str(input_path))
    _inject_null_relationship(input_path)

    findings = [
        Finding(
            rule_id="general-name-error",
            paragraph_index=0,
            line_number=1,
            original_text="中国人民银深圳市中心支行：",
            description="机构名缺字，应为'中国人民银行深圳市中心支行'",
            target_text="中国人民银",
        ),
    ]

    mark_errors_in_docx(input_path, output_path, findings)
    assert output_path.exists()
    marked = _marked_texts(output_path)
    assert "中国人民银" in marked


def test_mark_errors_marks_searchable_original_when_target_not_found(tmp_path: Path):
    """target_text 找不到时仍标原文，并在批注中给出人工定位文本."""
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _make_minimal_docx(input_path, ["原文里没有目标词。"])

    findings = [
        Finding(
            rule_id="general-typo",
            paragraph_index=0,
            line_number=1,
            original_text="原文里没有目标词。",
            description="'不存在'误写",
            target_text="不存在",
        ),
    ]

    mark_errors_in_docx(input_path, output_path, findings)
    assert output_path.exists()
    assert _marked_texts(output_path) == ["原文里没有目标词。"]


def test_mark_errors_relocates_by_exact_original_text_when_index_drifted(tmp_path: Path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _make_minimal_docx(input_path, ["第一段没有问题。", "第二段布署了工作。"])

    findings = [
        Finding(
            rule_id="general-typo",
            paragraph_index=0,
            line_number=1,
            original_text="第二段布署了工作。",
            description="'部署'误写为'布署'",
            target_text="布署",
        ),
    ]

    mark_errors_in_docx(input_path, output_path, findings)

    assert _marked_texts(output_path) == ["布署"]


def test_mark_errors_uses_long_content_anchor_before_claimed_paragraph(tmp_path: Path):
    input_path = tmp_path / "content_anchor.docx"
    output_path = tmp_path / "content_anchor_marked.docx"
    paragraphs = [
        "第一段也出现考字，但没有错误。",
        "中间说明。",
        "在示例流程中，建立综合考考量服务规模和风险水平的体系。",
    ]
    _make_minimal_docx(input_path, paragraphs)
    finding = Finding(
        rule_id="general-typo",
        paragraph_index=0,
        line_number=1,
        original_text=paragraphs[2],
        description="重复‘考’字，应为‘综合考量’",
        target_text="考考",
    )

    mark_errors_in_docx(input_path, output_path, [finding])

    assert _marked_offsets(output_path) == [
        (2, paragraphs[2].index("考考"), "考考")
    ]


def test_mark_errors_uses_locator_to_select_repeated_target(tmp_path: Path):
    input_path = tmp_path / "repeated_target.docx"
    output_path = tmp_path / "repeated_target_marked.docx"
    paragraph = "第一处使用英文逗号,不处理；真正的问题在同时,每期更新。"
    _make_minimal_docx(input_path, [paragraph])
    finding = Finding(
        rule_id="general-punctuation",
        paragraph_index=0,
        line_number=1,
        original_text="真正的问题在同时,每期更新。",
        description="中文句子中出现英文逗号",
        target_text=",",
    )

    mark_errors_in_docx(input_path, output_path, [finding])

    assert _marked_offsets(output_path) == [
        (0, paragraph.rindex(","), ",")
    ]


def test_mark_errors_fallback_still_adds_searchable_original_comment(tmp_path: Path):
    input_path = tmp_path / "fallback_comment.docx"
    output_path = tmp_path / "fallback_comment_marked.docx"
    _make_minimal_docx(input_path, ["当前段落内容已经发生变化。"])
    finding = Finding(
        rule_id="general-grammar",
        paragraph_index=0,
        line_number=1,
        original_text="原始错误句子将项目进度被纳入考核。",
        description="‘将……被纳入’句式杂糅",
        target_text="将项目进度被纳入考核",
    )

    mark_errors_in_docx(input_path, output_path, [finding])

    assert _marked_texts(output_path)
    from docx import Document

    comments = list(Document(str(output_path)).comments)
    assert len(comments) == 1
    assert "定位原文" in comments[0].text
    assert "原始错误句子将项目进度被纳入考核" in comments[0].text


def test_mark_errors_uses_same_reviewable_paragraph_index_as_parser(tmp_path: Path):
    input_path = tmp_path / "synthetic.docx"
    output_path = tmp_path / "synthetic_marked.docx"
    _make_minimal_docx(input_path, ["第一段正常。", "示例机购名称缺字："])
    parsed = parse_docx(input_path)
    target_index = parsed.paragraphs.index("示例机购名称缺字：")

    findings = [
        Finding(
            rule_id="general-name-error",
            paragraph_index=target_index,
            line_number=target_index + 1,
            original_text="示例机购名称缺字：",
            description="机构名错字，应为‘示例机构名称’",
            target_text="机购",
        ),
    ]

    mark_errors_in_docx(input_path, output_path, findings)

    marked = _marked_texts(output_path)
    assert "机购" in marked


def test_mark_errors_quote_pair_marks_exact_unmatched_quote(tmp_path: Path):
    input_path = tmp_path / "quote_input.docx"
    output_path = tmp_path / "quote_output.docx"
    text = (
        "甲方表示如需调整可“及时沟通”，"
        "同时说明“下一阶段将继续推进。"
    )
    _make_minimal_docx(input_path, [text])

    findings = check_quote_pair([text])

    mark_errors_in_docx(input_path, output_path, findings)

    marked = _marked_texts(output_path)
    assert marked == ["“"]


def test_build_comment_text_omits_rule_label():
    finding = Finding(
        rule_id="content-mismatch",
        paragraph_index=0,
        line_number=1,
        original_text="标题",
        description="标题为A，但正文为B，内容不匹配。",
        target_text="标题",
    )

    comment = _build_comment_text(1, finding)

    assert "【" not in comment
    assert "标题正文不匹配" not in comment
    assert "标题为A，但正文为B，内容不匹配。" in comment
    assert "定位原文：标题" in comment


def test_build_comment_text_hides_internal_paragraph_number():
    finding = Finding(
        rule_id="general-duplicate",
        paragraph_index=10,
        line_number=11,
        original_text="这是一段可以直接搜索的重复原文。",
        description="本段与段落369内容完全重复",
        target_text="重复原文",
    )

    comment = _build_comment_text(1, finding)

    assert "369" not in comment
    assert "文中另一处" in comment
    assert "定位原文" in comment
