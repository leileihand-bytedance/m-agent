"""存档机制 + 主流程逻辑的单元测试.

注:企微 Bot 联调(WSClient)需要真实凭证和网络,这里只测可测的逻辑:
  - .docx 后缀判断
  - 拒接非 .docx 文件
  - 存档到 data/reviews/<日期-序号>/
  - 大文件拒绝
"""

from __future__ import annotations

import asyncio
import json
import sys
import zipfile
from pathlib import Path

import pytest

# 让测试能 import app.*
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from app.review.main import (  # noqa: E402
    is_docx_filename,
    save_review,
    save_review_to_directory,
    ReviewConfig,
    _next_review_index,
    _prepare_review_reply_file,
    build_user_review_reply,
    _split_text_into_paragraphs,
    extract_text_content,
    _review_text,
    _start_neican_review,
    _resolve_text_registration_reply,
    _build_enter_welcome_text,
    _is_followup_review_request_text,
    _resolve_instruction_only_text_reply,
    _resolve_smalltalk_text_reply,
    _is_official_format_review_request_text,
    _configure_ws_client_timeouts,
    _build_delivery_failure_user_reply,
    _build_processing_failure_user_reply,
    _summarize_delivery_error,
    RecentSubmissionTracker,
    archive_multi_file_review,
    build_multi_file_review_reply,
    _build_file_ack,
    _extract_primary_inference_texts,
    _settle_auto_review_batch,
    _deliver_review_attachment,
    _build_queued_attachment_delivery,
    _run_review_task_worker_supervised,
    _reply_queued_review_acceptance,
    _send_queued_review_text,
    _process_queued_single_review,
    _queued_review_acceptance_message,
    _review_task_type_for_document_type,
    ReviewDeliveryStatusUncertain,
)
from app.review.document_type import DocumentType  # noqa: E402
from app.review.task_execution import (  # noqa: E402
    GENERAL_TEXT_REVIEW_TASK_TYPE,
    GENERAL_REVIEW_TASK_TYPE,
    HALF_MONTHLY_REVIEW_TASK_TYPE,
    NEICAN_REVIEW_TASK_TYPE,
    OFFICIAL_FORMAT_REVIEW_TASK_TYPE,
    GeneralReviewWorkspace,
)
from app.review.reviewer import ReviewResult, Finding  # noqa: E402
from app.review.multi_file_reviewer import (  # noqa: E402
    MultiFileReviewBundle,
    MultiFileReviewedDocument,
    MultiFileSource,
)
from app.platform.user_registry import RegistrationFlow, UserRegistry  # noqa: E402
from app.platform.task_status import write_task_status  # noqa: E402
from app.platform.attachment_delivery import DeliveryResult  # noqa: E402


# ============================================================
# 测试 1: .docx 后缀判断
# ============================================================

def test_is_docx_filename_accept():
    assert is_docx_filename("report.docx") is True
    assert is_docx_filename("Report.DOCX") is True
    assert is_docx_filename("汇报.docx") is True
    print("✅ test_is_docx_filename_accept: .docx 全部接受")


def test_is_docx_filename_reject():
    assert is_docx_filename("report.pdf") is False
    assert is_docx_filename("report.txt") is False
    assert is_docx_filename("report.md") is False
    assert is_docx_filename("report") is False
    assert is_docx_filename(None) is False
    assert is_docx_filename("") is False
    print("✅ test_is_docx_filename_reject: 非 .docx 全部拒")


def test_official_format_review_request_requires_explicit_format_wording():
    assert _is_official_format_review_request_text("帮我做一下格式审核") is True
    assert _is_official_format_review_request_text("按公文格式检查") is True
    assert _is_official_format_review_request_text("看看这个文件格式有没有问题") is True
    assert _is_official_format_review_request_text("帮我查一下格式") is True
    assert _is_official_format_review_request_text("帮我审一下这个材料") is False
    assert _is_official_format_review_request_text("请审核文字内容") is False
    assert _is_official_format_review_request_text("只审核内容，格式不用看") is False
    assert _is_official_format_review_request_text("内容要审，格式不用查") is False


# ============================================================
# 测试 2: 存档到 data/reviews/
# ============================================================


def test_archive_multi_file_review_creates_one_task_and_marked_outputs(tmp_path: Path):
    from docx import Document

    main_path = tmp_path / "正文.docx"
    attachment_path = tmp_path / "附件1.docx"
    for path, text in ((main_path, "详见附件1。"), (attachment_path, "附件1：名单")):
        document = Document()
        document.add_paragraph(text)
        document.save(path)
    finding = Finding(
        rule_id="multi-file-reference-missing",
        paragraph_index=0,
        line_number=1,
        original_text="详见附件1。",
        description="示例跨文件问题",
        target_text="附件1",
    )
    bundle = MultiFileReviewBundle(
        documents=(
            MultiFileReviewedDocument(
                source=MultiFileSource(0, "正文.docx", main_path, ("详见附件1。",)),
                doc_type=DocumentType.GENERAL,
                result=ReviewResult([finding], 10, 9, "正文.docx"),
            ),
            MultiFileReviewedDocument(
                source=MultiFileSource(1, "附件1.docx", attachment_path, ("附件1：名单",)),
                doc_type=DocumentType.GENERAL,
                result=ReviewResult([], 10, 10, "附件1.docx"),
            ),
        ),
        cross_file_finding_count=1,
        primary_file_index=0,
    )

    task_dir, marked_paths = archive_multi_file_review(
        reviews_dir=tmp_path / "reviews",
        sender="user-1",
        msgid="message-1",
        bundle=bundle,
    )

    assert len(list((task_dir / "input").glob("*.docx"))) == 2
    assert len(marked_paths) == 1
    assert marked_paths[0].exists()
    assert "正文.docx" in (task_dir / "output" / "report.md").read_text(encoding="utf-8")
    meta = json.loads((task_dir / "meta.json").read_text(encoding="utf-8"))
    assert meta["document_type"] == "multi_file"
    assert meta["cross_file_finding_count"] == 1
    assert meta["primary_file_index"] == 0
    assert meta["files"][0]["is_primary"] is True
    assert meta["files"][1]["is_primary"] is False
    status = json.loads((task_dir / "status.json").read_text(encoding="utf-8"))
    assert status["processing_status"] == "completed"
    assert status["delivery_status"] == "unknown"
    report = (task_dir / "output" / "report.md").read_text(encoding="utf-8")
    assert "主文件：正文.docx" in report


def test_build_multi_file_review_reply_only_summarizes_counts():
    bundle = MultiFileReviewBundle(
        documents=(
            MultiFileReviewedDocument(
                source=MultiFileSource(0, "正文.docx", Path("正文.docx"), ("正文",)),
                doc_type=DocumentType.GENERAL,
                result=ReviewResult([], 10, 10, "正文.docx"),
            ),
            MultiFileReviewedDocument(
                source=MultiFileSource(1, "附件1.docx", Path("附件1.docx"), ("附件",)),
                doc_type=DocumentType.GENERAL,
                result=ReviewResult([], 10, 10, "附件1.docx"),
            ),
        ),
        cross_file_finding_count=0,
        primary_file_index=0,
    )

    reply = build_multi_file_review_reply(bundle, marked_file_count=0)

    assert "2 份文件" in reply
    assert "主文件：正文.docx" in reply
    assert "正文.docx：0 处" in reply
    assert "跨文件问题：0 处" in reply
    assert "错误1" not in reply


def test_build_multi_file_review_reply_does_not_claim_files_were_already_sent():
    bundle = MultiFileReviewBundle(
        documents=(
            MultiFileReviewedDocument(
                source=MultiFileSource(0, "正文.docx", Path("正文.docx"), ("正文",)),
                doc_type=DocumentType.GENERAL,
                result=ReviewResult([], 10, 10, "正文.docx"),
            ),
            MultiFileReviewedDocument(
                source=MultiFileSource(1, "附件1.docx", Path("附件1.docx"), ("附件",)),
                doc_type=DocumentType.GENERAL,
                result=ReviewResult([], 10, 10, "附件1.docx"),
            ),
        ),
        cross_file_finding_count=1,
        primary_file_index=0,
    )

    reply = build_multi_file_review_reply(bundle, marked_file_count=2)

    assert "共生成 2 份带批注的文档，将继续发送" in reply
    assert "已返回" not in reply

def _make_fake_docx_bytes() -> bytes:
    """构造一个最小的 .docx 文件,返回字节."""
    import io
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("word/document.xml", '''<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:r><w:t>测试内容</w:t></w:r></w:p>
  </w:body>
</w:document>''')
    return buf.getvalue()


def test_save_review_creates_directory():
    """存档应创建标准目录结构."""
    import tempfile
    import shutil

    with tempfile.TemporaryDirectory() as tmpdir:
        reviews_dir = Path(tmpdir) / "reviews"
        reviews_dir.mkdir()  # 提前创建
        file_bytes = _make_fake_docx_bytes()

        # 构造一个假审核结果
        result = ReviewResult(
            findings=[
                Finding(
                    rule_id="dupe-char-check",
                    paragraph_index=0,
                    line_number=1,
                    original_text="我们要的的",
                    description="重复字",
                ),
            ],
            total_rules=10,
            passed_rules=9,
            filename="test.docx",
        )

        review_dir = save_review(
            reviews_dir=reviews_dir,
            file_bytes=file_bytes,
            original_filename="测试报告.docx",
            sender="user123",
            msgid="msg456",
            result=result,
            parsed_paragraphs=["测试内容", "第二段"],
        )

        # 验证目录结构
        assert review_dir.exists()
        assert (review_dir / "input").exists()
        assert (review_dir / "output").exists()
        assert (review_dir / "input" / "测试报告.docx").exists()
        assert (review_dir / "output" / "report.md").exists()
        assert (review_dir / "meta.json").exists()
        status = json.loads((review_dir / "status.json").read_text(encoding="utf-8"))
        assert status["processing_status"] == "completed"
        assert status["delivery_status"] == "unknown"

        # 验证 report.md 内容
        report = (review_dir / "output" / "report.md").read_text(encoding="utf-8")
        assert "测试报告.docx" in report
        assert "dupe-char-check" in report

        # 验证 meta.md 内容
        meta = (review_dir / "meta.json").read_text(encoding="utf-8")
        assert "user123" in meta
        assert "msg456" in meta

        print(f"✅ test_save_review_creates_directory: 存档目录 {review_dir.name} 结构正确")


def test_save_review_to_existing_queue_directory_keeps_running_status(tmp_path: Path):
    review_dir = tmp_path / "reviews" / "2026" / "07" / "queued-task"
    write_task_status(
        review_dir,
        processing_status="running",
        delivery_status="unknown",
        source="task_execution",
        state_version=2,
    )
    (review_dir / "meta.json").write_text(
        json.dumps(
            {
                "task_id": "queue-task-1",
                "queue_mode": "persistent",
                "task_type": GENERAL_REVIEW_TASK_TYPE,
                "queued_at": "2026-07-16T09:00:00+08:00",
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    result = ReviewResult(
        findings=[],
        total_rules=10,
        passed_rules=10,
        filename="材料.docx",
    )

    saved = save_review_to_directory(
        review_dir=review_dir,
        file_bytes=_make_fake_docx_bytes(),
        original_filename="材料.docx",
        sender="user-1",
        msgid="message-1",
        task_id="queue-task-1",
        result=result,
        parsed_paragraphs=["测试内容"],
        doc_type=DocumentType.GENERAL,
        mark_processing_completed=False,
    )

    assert saved == review_dir
    assert (review_dir / "input" / "材料.docx").exists()
    assert (review_dir / "output" / "report.md").exists()
    status = json.loads((review_dir / "status.json").read_text(encoding="utf-8"))
    assert status["processing_status"] == "running"
    assert status["state_version"] == 2
    meta = json.loads((review_dir / "meta.json").read_text(encoding="utf-8"))
    assert meta["task_id"] == "queue-task-1"
    assert meta["message_id"] == "message-1"
    assert meta["queue_mode"] == "persistent"
    assert meta["task_type"] == GENERAL_REVIEW_TASK_TYPE
    assert meta["queued_at"] == "2026-07-16T09:00:00+08:00"


@pytest.mark.parametrize(
    ("doc_type", "task_type"),
    [
        (DocumentType.GENERAL, GENERAL_REVIEW_TASK_TYPE),
        (DocumentType.HALF_MONTHLY, HALF_MONTHLY_REVIEW_TASK_TYPE),
        (DocumentType.NEI_CAN, NEICAN_REVIEW_TASK_TYPE),
        (DocumentType.OFFICIAL_FORMAT, OFFICIAL_FORMAT_REVIEW_TASK_TYPE),
    ],
)
def test_review_task_type_mapping_covers_every_single_file_review(
    doc_type: DocumentType,
    task_type: str,
):
    assert _review_task_type_for_document_type(doc_type) == task_type


def test_queued_review_acceptance_message_names_actual_review_type():
    created_message = _queued_review_acceptance_message(
        review_label="半月报审核",
        created=True,
        input_label="这份文件",
    )
    duplicate_message = _queued_review_acceptance_message(
        review_label="文字审核",
        created=False,
        input_label="这段文字",
    )

    assert created_message == "收到，正在进行半月报审核，完成后会自动发送结果。"
    assert duplicate_message == "这段文字已经在处理中，无需重复提交。完成后会自动发送结果。"


@pytest.mark.anyio
async def test_file_ack_suppresses_redundant_queue_acceptance():
    class FakeWsClient:
        def __init__(self) -> None:
            self.replies = []

        async def reply_stream(self, frame, stream_id, message, finish):
            self.replies.append((frame, stream_id, message, finish))

    ws_client = FakeWsClient()

    sent = await _reply_queued_review_acceptance(
        ws_client,
        {"msgid": "message-1"},
        "review-queued-1",
        review_label="通用审核",
        created=True,
        input_label="这份文件",
        acknowledgment_already_sent=True,
    )

    assert sent is False
    assert ws_client.replies == []


@pytest.mark.parametrize(
    ("task_type", "filename", "engine_name"),
    [
        (GENERAL_REVIEW_TASK_TYPE, "材料.docx", "general"),
        (HALF_MONTHLY_REVIEW_TASK_TYPE, "信息动态半月报.docx", "halfmonthly"),
        (NEICAN_REVIEW_TASK_TYPE, "信息内参周报.docx", "neican"),
        (OFFICIAL_FORMAT_REVIEW_TASK_TYPE, "请示.docx", "official_format"),
    ],
)
def test_persistent_single_review_routes_to_existing_engine(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    task_type: str,
    filename: str,
    engine_name: str,
):
    from app.review import general_reviewer, halfmonthly_reviewer, official_format_checker, reviewer

    calls: list[str] = []

    def passed_result(name: str) -> ReviewResult:
        return ReviewResult(findings=[], total_rules=1, passed_rules=1, filename=name)

    async def fake_general(_paragraphs, _rules, name, **kwargs):
        calls.append("general")
        assert kwargs["profile"].profile_id == "general_docx"
        return passed_result(name)

    async def fake_halfmonthly(_paragraphs, _rules, name, **_kwargs):
        calls.append("halfmonthly")
        return passed_result(name)

    async def fake_phase1(_paragraphs, _rules, name, **_kwargs):
        calls.append("neican-phase1")
        return passed_result(name)

    async def fake_phase2(_paragraphs, _rules, name, **_kwargs):
        calls.append("neican-phase2")
        return passed_result(name)

    def fake_official_format(_path, name):
        calls.append("official_format")
        return passed_result(name)

    monkeypatch.setattr(general_reviewer, "review_general", fake_general)
    monkeypatch.setattr(halfmonthly_reviewer, "review_halfmonthly", fake_halfmonthly)
    monkeypatch.setattr(reviewer, "review_phase1", fake_phase1)
    monkeypatch.setattr(reviewer, "review_phase2", fake_phase2)
    monkeypatch.setattr(official_format_checker, "review_official_format", fake_official_format)

    task_dir = tmp_path / "reviews" / "queued-task"
    input_dir = task_dir / "input"
    (task_dir / "output").mkdir(parents=True)
    input_dir.mkdir(parents=True)
    input_file = input_dir / filename
    from docx import Document

    document = Document()
    document.add_paragraph("测试内容")
    document.save(input_file)
    workspace = GeneralReviewWorkspace(
        task_id="task-123",
        task_dir=task_dir,
        input_file=input_file,
        filename=filename,
        sender_userid="user-1",
        sender_name="User One",
        task_type=task_type,
    )
    config = ReviewConfig(
        wecom_bot_id="bot",
        wecom_bot_secret="secret",
        rules_path=tmp_path / "rules.md",
        reviews_dir=tmp_path / "reviews",
        logs_dir=tmp_path / "logs",
        admin_user_id="",
        admin_name="",
        notification_cooldown=300,
        direct_admin_notifications=False,
        require_registration=False,
    )

    delivery = asyncio.run(
        _process_queued_single_review(
            workspace,
            config=config,
            neican_rules_text="rules",
        )
    )

    assert delivery.kind == "text"
    if engine_name == "neican":
        assert sorted(calls) == ["neican-phase1", "neican-phase2"]
    else:
        assert calls == [engine_name]
    assert (task_dir / "output" / "report.md").is_file()


def test_persistent_text_review_uses_snapshotted_text_and_returns_text(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    import app.review.main as review_main

    seen: list[str] = []

    async def fake_text_review(content: str, _config: ReviewConfig):
        seen.append(content)
        return (
            ReviewResult(
                findings=[],
                total_rules=1,
                passed_rules=1,
                filename="文字消息",
            ),
            [content],
        )

    monkeypatch.setattr(review_main, "_review_text_result", fake_text_review)
    task_dir = tmp_path / "reviews" / "queued-text"
    input_dir = task_dir / "input"
    (task_dir / "output").mkdir(parents=True)
    input_dir.mkdir(parents=True)
    input_file = input_dir / "文字消息.txt"
    input_file.write_text("需要审核的文字", encoding="utf-8")
    workspace = GeneralReviewWorkspace(
        task_id="task-text",
        task_dir=task_dir,
        input_file=input_file,
        filename="文字消息.txt",
        sender_userid="user-1",
        sender_name="User One",
        task_type=GENERAL_TEXT_REVIEW_TASK_TYPE,
        input_kind="text",
    )
    config = ReviewConfig(
        wecom_bot_id="bot",
        wecom_bot_secret="secret",
        rules_path=tmp_path / "rules.md",
        reviews_dir=tmp_path / "reviews",
        logs_dir=tmp_path / "logs",
        admin_user_id="",
        admin_name="",
        notification_cooldown=300,
        direct_admin_notifications=False,
        require_registration=False,
    )

    delivery = asyncio.run(
        _process_queued_single_review(
            workspace,
            config=config,
            neican_rules_text="rules",
        )
    )

    assert seen == ["需要审核的文字"]
    assert delivery.kind == "text"
    assert "未发现低级错误" in delivery.text
    assert (task_dir / "output" / "report.md").is_file()


def test_save_review_increments_index():
    """同一日多次审核,序号应递增."""
    import tempfile

    with tempfile.TemporaryDirectory() as tmpdir:
        reviews_dir = Path(tmpdir) / "reviews"
        reviews_dir.mkdir()

        # 存档 3 次
        dirs = []
        for i in range(3):
            result = ReviewResult(
                findings=[], total_rules=10, passed_rules=10, filename=f"f{i}.docx"
            )
            d = save_review(
                reviews_dir=reviews_dir,
                file_bytes=b"x",
                original_filename=f"f{i}.docx",
                sender="u",
                msgid=f"m{i}",
                result=result,
                parsed_paragraphs=[],
            )
            dirs.append(d)

        # 验证序号递增
        names = [d.name for d in dirs]
        # 名字形如 "20260613-001" "20260613-002" "20260613-003"
        suffixes = [n.split("-")[-1] for n in names]
        assert suffixes == ["001", "002", "003"], f"序号应为 001/002/003,实际 {suffixes}"
        print(f"✅ test_save_review_increments_index: 3 次存档序号 {suffixes}")


def test_save_review_sanitizes_filename():
    """特殊字符文件名应被清洗."""
    import tempfile

    with tempfile.TemporaryDirectory() as tmpdir:
        reviews_dir = Path(tmpdir) / "reviews"
        reviews_dir.mkdir()
        result = ReviewResult(
            findings=[], total_rules=10, passed_rules=10, filename="x"
        )

        # 含特殊字符的文件名
        review_dir = save_review(
            reviews_dir=reviews_dir,
            file_bytes=b"x",
            original_filename="../../etc/passwd.docx",  # 路径穿越
            sender="u",
            msgid="m",
            result=result,
            parsed_paragraphs=[],
        )

        # 验证文件确实在 source 下,没有逃逸
        source_files = list((review_dir / "input").iterdir())
        assert len(source_files) == 1
        saved = source_files[0]
        assert "passwd" in saved.name
        # 不应包含 .. 或 /
        assert ".." not in saved.name
        assert "/" not in saved.name
        print(f"✅ test_save_review_sanitizes_filename: 危险文件名被清洗为 {saved.name}")


def test_save_review_supports_text_message_archive(tmp_path: Path):
    reviews_dir = tmp_path / "reviews"
    reviews_dir.mkdir()

    result = ReviewResult(
        findings=[],
        total_rules=10,
        passed_rules=10,
        filename="文字消息",
    )

    review_dir = save_review(
        reviews_dir=reviews_dir,
        file_bytes=None,
        original_filename="文字消息.txt",
        sender="user123",
        msgid="msg-text-001",
        result=result,
        parsed_paragraphs=["第一段文字。", "第二段文字。"],
        text_content="第一段文字。\n\n第二段文字。",
        doc_type=DocumentType.GENERAL,
    )

    text_source = review_dir / "input" / "文字消息.txt"
    assert text_source.exists()
    assert text_source.read_text(encoding="utf-8") == "第一段文字。\n\n第二段文字。"


def test_prepare_review_reply_file_builds_marked_doc_for_neican_findings(tmp_path: Path):
    from docx import Document

    source_doc = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph("国务院政策例行吹风会介绍推进城市更新工作有关情况")
    doc.save(str(source_doc))

    review_dir = tmp_path / "review"
    source_dir = review_dir / "input"
    source_dir.mkdir(parents=True)
    saved_source = source_dir / "内参周报.docx"
    saved_source.write_bytes(source_doc.read_bytes())

    findings = [
        Finding(
            rule_id="content-mismatch",
            paragraph_index=0,
            line_number=1,
            original_text="国务院政策例行吹风会介绍推进城市更新工作有关情况",
            description="标题和正文讲的不是同一件事",
            target_text="国务院政策例行吹风会介绍推进城市更新工作有关情况",
        )
    ]

    reply_file = _prepare_review_reply_file(review_dir, "内参周报.docx", findings)

    assert reply_file is not None
    assert reply_file.exists()
    assert reply_file.name == "marked_内参周报.docx"


def test_prepare_review_reply_file_returns_source_when_no_findings(tmp_path: Path):
    source_doc = tmp_path / "source.docx"
    source_doc.write_bytes(b"fake-docx")

    review_dir = tmp_path / "review"
    source_dir = review_dir / "input"
    source_dir.mkdir(parents=True)
    saved_source = source_dir / "内参周报.docx"
    saved_source.write_bytes(source_doc.read_bytes())

    reply_file = _prepare_review_reply_file(review_dir, "内参周报.docx", [])

    assert reply_file is None


# ============================================================
# 测试 3: 配置加载
# ============================================================

def test_load_config_from_env_file():
    """从 .env 加载配置."""
    import tempfile
    import os

    with tempfile.TemporaryDirectory() as tmpdir:
        env_path = Path(tmpdir) / ".env"
        env_path.write_text(
            "\n".join([
                "WECOM_REVIEW_BOT_ID=test_bot_id_12345",
                "WECOM_REVIEW_BOT_SECRET=test_secret_abcde",
                "M_AGENT_REVIEW_RULES=app/data/rules.md",
                "M_AGENT_REVIEWS_DIR=data/reviews",
                f"M_AGENT_DATA_DIR={Path(tmpdir) / 'M-Agent-Files'}",
                "M_AGENT_LOG_MAX_MB=8",
                "M_AGENT_REVIEW_AUTO_BATCH_SECONDS=6.5",
                "REVIEW_REPLY_ACK_TIMEOUT_SECONDS=45",
            ]),
            encoding="utf-8",
        )

        # 临时切到 tmpdir,因为 load_config 用的是 _ROOT
        # 这里直接调 load_config 传 env_path
        from app.review.main import load_config
        config = load_config(env_path)

        assert config.wecom_bot_id == "test_bot_id_12345"
        assert config.wecom_bot_secret == "test_secret_abcde"
        assert config.rules_path.name == "rules.md"
        assert config.reviews_dir.name == "reviews"
        assert config.user_registry_path == (
            Path(tmpdir) / "M-Agent-Files" / "runtime" / "users" / "review_users.yaml"
        ).resolve()
        assert config.max_file_size_mb == 10  # 默认值
        assert config.reply_ack_timeout_seconds == 45.0
        assert config.log_max_bytes == 8 * 1024 * 1024
        assert config.direct_admin_notifications is False
        assert config.auto_batch_seconds == 6.5
        print("✅ test_load_config_from_env_file: 配置加载正确")


def test_review_load_config_uses_single_external_data_root(tmp_path: Path):
    data_root = tmp_path / "M-Agent-Files"
    env_path = tmp_path / ".env"
    env_path.write_text(
        "\n".join(
            [
                "WECOM_REVIEW_BOT_ID=bot-id",
                "WECOM_REVIEW_BOT_SECRET=bot-secret",
                f"M_AGENT_DATA_DIR={data_root}",
            ]
        ),
        encoding="utf-8",
    )

    from app.review.main import load_config

    config = load_config(env_path)

    assert config.reviews_dir == data_root / "tasks" / "review"
    assert config.logs_dir == data_root / "runtime" / "logs"
    assert config.ops_events_dir == data_root / "runtime" / "ops" / "events"
    assert config.ops_heartbeat_dir == data_root / "runtime" / "ops" / "heartbeats"
    assert config.user_registry_path == data_root / "runtime" / "users" / "review_users.yaml"
    assert config.intake_dir == data_root / "runtime" / "intake" / "review"
    assert config.intake_ttl_seconds == 1800
    assert config.auto_batch_seconds == 8.0
    assert config.task_queue_db == data_root / "runtime" / "task-execution" / "review.sqlite3"
    assert config.task_worker_count == 1
    assert config.task_poll_seconds == 0.25
    assert config.task_recovery_seconds == 5.0
    assert config.task_lease_seconds == 120


def test_default_file_ack_keeps_original_review_time_expectation():
    reply = _build_file_ack(None)

    assert reply == (
        "收到文件啦，正在加紧审核，请稍等"
        "（模型反应有点慢，你可以先干点别的，一会儿再来看）……"
    )


def test_extract_primary_inference_texts_reads_each_real_docx(tmp_path: Path):
    from docx import Document
    from app.platform.models import UploadedFile

    main_path = tmp_path / "材料甲.docx"
    attachment_path = tmp_path / "材料乙.docx"
    for path, paragraphs in (
        (main_path, ["会议通知", "请填写附件1。"]),
        (attachment_path, ["附件1", "议案意见反馈表"]),
    ):
        document = Document()
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        document.save(path)

    texts = _extract_primary_inference_texts(
        (
            UploadedFile(filename=main_path.name, stored_path=str(main_path)),
            UploadedFile(filename=attachment_path.name, stored_path=str(attachment_path)),
        )
    )

    assert "请填写附件1" in texts[0]
    assert texts[1].startswith("附件1")


def test_settle_auto_review_batch_routes_multiple_files_without_user_text(tmp_path: Path):
    import asyncio
    from docx import Document
    from app.platform.models import UploadedFile
    from app.review.intake import ReviewIntakeStore

    files: list[UploadedFile] = []
    for filename, paragraphs in (
        ("材料甲.docx", ["附件1", "议案意见反馈表"]),
        ("材料乙.docx", ["会议通知", "请填写附件1。"]),
    ):
        path = tmp_path / filename
        document = Document()
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        document.save(path)
        files.append(UploadedFile(filename=filename, content=path.read_bytes()))

    store = ReviewIntakeStore(storage_dir=tmp_path / "intake")
    store.add_file(channel="wecom", sender_userid="user-1", file=files[0])
    queued = store.add_file(channel="wecom", sender_userid="user-1", file=files[1])

    decision = asyncio.run(
        _settle_auto_review_batch(
            store,
            channel="wecom",
            sender_userid="user-1",
            expected_revision=queued.revision,
            delay_seconds=0,
        )
    )

    assert decision.action == "run_multi"
    assert decision.primary_file_index == 1


def test_configure_ws_client_timeouts_overrides_sdk_reply_ack_timeout():
    class FakeManager:
        def __init__(self) -> None:
            self._reply_ack_timeout = 5.0

    class FakeClient:
        def __init__(self) -> None:
            self._ws_manager = FakeManager()

    client = FakeClient()

    _configure_ws_client_timeouts(client, reply_ack_timeout_seconds=30.0)

    assert client._ws_manager._reply_ack_timeout == 30.0


def test_build_delivery_failure_user_reply_hides_req_id_detail():
    exc = RuntimeError("Reply ack timeout (5.0s) for reqId: secret-req-id")

    reply = _build_delivery_failure_user_reply("标注文档", exc)

    assert "审核已经完成" in reply
    assert "标注文档发送失败" in reply
    assert "企业微信上传回执超时" in reply
    assert "管理员" in reply
    assert "secret-req-id" not in reply


def test_summarize_delivery_error_maps_upload_failure():
    exc = RuntimeError("Upload failed: 3 chunk(s) failed")

    assert _summarize_delivery_error(exc) == "企业微信文件上传失败"


def test_review_attachment_delivery_uses_public_status_and_metrics(tmp_path: Path):
    import asyncio

    from app.platform.attachment_delivery import AttachmentDelivery
    from app.platform.task_status import write_task_status

    class FakeWsClient:
        def __init__(self):
            self.uploaded = []
            self.replied = []

        async def upload_media(self, content, *, type, filename):
            self.uploaded.append((content, type, filename))
            return {"media_id": "media-1"}

        async def reply_media(self, frame, media_type, media_id):
            self.replied.append((frame, media_type, media_id))

        async def reply_stream(self, frame, req_id, text, finish):
            raise AssertionError("成功交付不应发送失败提示")

    task_dir = tmp_path / "review-task"
    output = task_dir / "output" / "marked_材料.docx"
    output.parent.mkdir(parents=True)
    output.write_bytes(b"marked-docx")
    write_task_status(task_dir, processing_status="completed", delivery_status="unknown")
    ws_client = FakeWsClient()

    delivered = asyncio.run(
        _deliver_review_attachment(
            attachment_delivery=AttachmentDelivery(),
            ws_client=ws_client,
            frame={"msgid": "msg-1"},
            path=output,
            sender="user-1",
            sender_name="test-user",
            label="标注文档",
            req_id_factory=lambda prefix: f"{prefix}-1",
        )
    )

    assert delivered is True
    assert ws_client.uploaded == [(b"marked-docx", "file", "marked_材料.docx")]
    assert (task_dir / "delivery.json").is_file()
    status = json.loads((task_dir / "status.json").read_text(encoding="utf-8"))
    assert status["delivery_status"] == "delivered"


def test_queued_review_attachment_failure_is_reported_only_by_task_service(
    tmp_path: Path,
):
    import asyncio

    class FailedDelivery:
        async def deliver(self, *, ws_client, request):
            return DeliveryResult(
                delivered=False,
                status="failed",
                attempts=3,
                size_bytes=10,
                estimated_chunks=1,
                upload_elapsed_seconds=1.0,
                error_code="upload_failed",
                user_message="附件发送失败",
                metrics_path=None,
            )

    class FakeWsClient:
        def __init__(self) -> None:
            self.messages = []

        async def send_message(self, chat_id, body):
            self.messages.append((chat_id, body))

    task_dir = tmp_path / "review-task"
    output = task_dir / "output" / "marked_材料.docx"
    output.parent.mkdir(parents=True)
    output.write_bytes(b"marked-docx")
    ws_client = FakeWsClient()

    delivered = asyncio.run(
        _deliver_review_attachment(
            attachment_delivery=FailedDelivery(),
            ws_client=ws_client,
            frame=None,
            chat_id="user-1",
            path=output,
            sender="user-1",
            sender_name="test-user",
            label="标注文档",
            req_id_factory=lambda prefix: f"{prefix}-1",
            notify_user_on_failure=False,
        )
    )

    assert delivered is False
    assert ws_client.messages == []


def test_queued_review_text_timeout_is_not_retried():
    import asyncio

    class FakeWsClient:
        def __init__(self) -> None:
            self.calls = 0

        async def send_message(self, chat_id, body):
            self.calls += 1
            raise asyncio.TimeoutError

    ws_client = FakeWsClient()

    with pytest.raises(asyncio.TimeoutError):
        asyncio.run(
            _send_queued_review_text(
                ws_client,
                "user-1",
                "审核完成",
                timeout_seconds=0.01,
            )
        )

    assert ws_client.calls == 1


def test_queued_review_text_uses_sdk_supported_markdown_message():
    import asyncio

    class FakeWsClient:
        def __init__(self) -> None:
            self.messages = []

        async def send_message(self, chat_id, body):
            self.messages.append((chat_id, body))

    ws_client = FakeWsClient()

    sent = asyncio.run(_send_queued_review_text(ws_client, "user-1", "审核完成"))

    assert sent is True
    assert ws_client.messages == [
        (
            "user-1",
            {"msgtype": "markdown", "markdown": {"content": "审核完成"}},
        )
    ]


def test_queued_review_attachment_reply_timeout_is_not_retried(tmp_path: Path):
    import asyncio

    class FakeWsClient:
        def __init__(self) -> None:
            self.upload_calls = 0
            self.send_calls = 0

        async def upload_media(self, content, *, type, filename):
            self.upload_calls += 1
            return {"media_id": "media-1"}

        async def send_media_message(self, chat_id, media_type, media_id):
            self.send_calls += 1
            raise asyncio.TimeoutError

    task_dir = tmp_path / "review-task"
    output = task_dir / "output" / "marked_材料.docx"
    output.parent.mkdir(parents=True)
    output.write_bytes(b"marked-docx")
    ws_client = FakeWsClient()

    with pytest.raises(ReviewDeliveryStatusUncertain):
        asyncio.run(
            _deliver_review_attachment(
                attachment_delivery=_build_queued_attachment_delivery(None),
                ws_client=ws_client,
                frame=None,
                chat_id="user-1",
                path=output,
                sender="user-1",
                sender_name="test-user",
                label="标注文档",
                req_id_factory=lambda prefix: f"{prefix}-1",
                notify_user_on_failure=False,
                raise_on_uncertain=True,
            )
        )

    assert ws_client.upload_calls == 1
    assert ws_client.send_calls == 1


def test_review_task_worker_supervisor_records_failure_and_restarts():
    import asyncio

    class FakeExecutor:
        def __init__(self) -> None:
            self.calls = 0

        async def run_forever(self, *, stop_event, **kwargs):
            self.calls += 1
            if self.calls == 1:
                raise RuntimeError("worker stopped")
            stop_event.set()

    class FakeOpsLogger:
        def __init__(self) -> None:
            self.events = []

        def record(self, **payload):
            self.events.append(payload)

    executor = FakeExecutor()
    ops_logger = FakeOpsLogger()

    asyncio.run(
        _run_review_task_worker_supervised(
            task_executor=executor,
            stop_event=asyncio.Event(),
            poll_interval=0.01,
            worker_count=1,
            recovery_interval=0.01,
            restart_delay_seconds=0,
            ops_event_logger=ops_logger,
        )
    )

    assert executor.calls == 2
    assert len(ops_logger.events) == 1
    assert ops_logger.events[0]["subject"] == "审核后台任务 worker 异常退出"


def test_load_config_missing_required():
    """缺少必填字段应报错."""
    import tempfile

    with tempfile.TemporaryDirectory() as tmpdir:
        env_path = Path(tmpdir) / ".env"
        env_path.write_text("# 空配置", encoding="utf-8")

        from app.review.main import load_config
        try:
            load_config(env_path)
            assert False, "应该抛 ValueError"
        except ValueError as e:
            assert "WECOM_REVIEW_BOT_ID" in str(e)
            print("✅ test_load_config_missing_required: 缺字段时正确报错")


# ============================================================
# 测试 4: 文字消息审核
# ============================================================

def test_split_text_into_paragraphs():
    """文本按空行/换行拆分段落,过滤空段."""
    assert _split_text_into_paragraphs("第一段。\n\n第二段。") == ["第一段。", "第二段。"]
    assert _split_text_into_paragraphs("A\nB\nC") == ["A", "B", "C"]
    assert _split_text_into_paragraphs("  \n\n  ") == []


def test_split_text_into_paragraphs_keeps_attachment_lines_after_blank_line():
    text = (
        "请填写附件1议案意见反馈表。\n\n"
        "附件 1：换届工作领导小组名单\n"
        "附件 7：议案意见反馈表"
    )

    assert _split_text_into_paragraphs(text) == [
        "请填写附件1议案意见反馈表。",
        "附件 1：换届工作领导小组名单",
        "附件 7：议案意见反馈表",
    ]


def test_extract_text_content():
    """从企微文本消息 frame 中提取 content."""
    frame = {
        "body": {
            "text": {"content": "  测试文字  "},
        }
    }
    assert extract_text_content(frame) == "测试文字"

    assert extract_text_content({}) is None
    assert extract_text_content({"body": {}}) is None
    assert extract_text_content({"body": {"text": {}}}) is None


def test_resolve_text_registration_reply_registers_valid_name(tmp_path: Path):
    registry = UserRegistry(tmp_path / "users.yaml")
    flow = RegistrationFlow(registry, require_registration=True)

    handled, reply = _resolve_text_registration_reply(flow, "new_user", "Tom")

    assert handled is True
    assert reply == (
        "你好，Tom：\n"
        "我可以帮你审内参、半月报，或者其他文字材料，"
        "直接发文字、docx、html或pptx给我就可以。"
        "另外请注意，涉及行内数据请务必脱敏哦。"
    )
    assert registry.get_name("new_user") == "Tom"


def test_resolve_text_registration_reply_rejects_invalid_name(tmp_path: Path):
    registry = UserRegistry(tmp_path / "users.yaml")
    flow = RegistrationFlow(registry, require_registration=True)

    handled, reply = _resolve_text_registration_reply(flow, "new_user", "Tom@123")

    assert handled is True
    assert "请发送一个有效的名字" in reply
    assert not registry.is_registered("new_user")


def test_resolve_text_registration_reply_prompts_again_for_common_greeting(tmp_path: Path):
    registry = UserRegistry(tmp_path / "users.yaml")
    flow = RegistrationFlow(registry, require_registration=True)

    handled, reply = _resolve_text_registration_reply(flow, "new_user", "你好")

    assert handled is True
    assert reply == flow.ask_name_message()
    assert not registry.is_registered("new_user")


def test_resolve_text_registration_reply_prompts_again_for_review_request_text(tmp_path: Path):
    registry = UserRegistry(tmp_path / "users.yaml")
    flow = RegistrationFlow(registry, require_registration=True)

    handled, reply = _resolve_text_registration_reply(flow, "new_user", "我要审个材料")

    assert handled is True
    assert reply == flow.ask_name_message()
    assert not registry.is_registered("new_user")


def test_build_enter_welcome_text_asks_name_for_new_user(tmp_path: Path):
    registry = UserRegistry(tmp_path / "users.yaml")
    flow = RegistrationFlow(registry, require_registration=True)

    reply = _build_enter_welcome_text(flow, "new_user")

    assert reply == (
        "欢迎使用智能审核BOT！\n"
        "这是你第一次使用，先互相认识一下吧，请先告诉我你的英文名（例如：Jack）。"
    )


def test_build_enter_welcome_text_keeps_regular_welcome_for_registered_user(tmp_path: Path):
    registry = UserRegistry(tmp_path / "users.yaml")
    registry.register("old_user", "Alice")
    flow = RegistrationFlow(registry, require_registration=True)

    reply = _build_enter_welcome_text(flow, "old_user")

    assert reply == (
        "你好，需要我帮你审核什么呢？请直接发送 .docx、.html/.htm、.pptx 文件"
        "或直接发送文字，我会认真审核。"
    )


def test_is_followup_review_request_text_matches_short_prompt():
    assert _is_followup_review_request_text("帮我审一下") is True
    assert _is_followup_review_request_text("帮我看看有无问题") is True
    assert _is_followup_review_request_text("请审核一下。") is True
    assert _is_followup_review_request_text("帮我审一下这个材料") is True
    assert _is_followup_review_request_text("也做一下文字审核") is True
    assert _is_followup_review_request_text("再做一下内容审核") is True


def test_is_followup_review_request_text_does_not_match_real_content():
    assert _is_followup_review_request_text("这是需要审核的正文第一段内容") is False
    assert _is_followup_review_request_text("今天召开专题会议，研究下半年工作安排。") is False
    assert _is_followup_review_request_text("文字审核工作应重点关注材料中的事实错误。") is False


def test_recent_submission_tracker_ignores_followup_prompt_after_file():
    tracker = RecentSubmissionTracker(ttl_seconds=90)
    tracker.remember("u1", "file", now=100.0)

    assert tracker.should_ignore_text_review("u1", "帮我审一下", now=120.0) is True


def test_recent_submission_tracker_does_not_ignore_followup_without_recent_submission():
    tracker = RecentSubmissionTracker(ttl_seconds=90)

    assert tracker.should_ignore_text_review("u1", "帮我审一下", now=120.0) is False


def test_recent_submission_tracker_does_not_ignore_real_text_after_file():
    tracker = RecentSubmissionTracker(ttl_seconds=90)
    tracker.remember("u1", "file", now=100.0)

    assert tracker.should_ignore_text_review("u1", "第一段：今天召开专题会议。", now=120.0) is False


def test_resolve_instruction_only_text_reply_prompts_user_to_send_material_first():
    tracker = RecentSubmissionTracker(ttl_seconds=90)

    reply = _resolve_instruction_only_text_reply(tracker, "u1", "帮我审一下这个材料", now=100.0)

    assert reply == "收到，请把需要审核的文字、.docx、.html/.htm或.pptx文件发给我，我来帮你看。"


def test_resolve_instruction_only_text_reply_treats_followup_as_continue_when_recent_submission_exists():
    tracker = RecentSubmissionTracker(ttl_seconds=90)
    tracker.remember("u1", "file", now=100.0)

    reply = _resolve_instruction_only_text_reply(tracker, "u1", "帮我审一下这个材料", now=120.0)

    assert reply == "收到，我会按你刚发的内容继续审核，请稍等……"


def test_resolve_instruction_only_text_reply_does_not_audit_content_review_command():
    tracker = RecentSubmissionTracker(ttl_seconds=90)
    tracker.remember("u1", "file", now=100.0)

    reply = _resolve_instruction_only_text_reply(
        tracker,
        "u1",
        "也做一下文字审核",
        now=120.0,
    )

    assert reply == "收到，我会按你刚发的内容继续审核，请稍等……"


def test_resolve_smalltalk_text_reply_handles_ack_text():
    assert _resolve_smalltalk_text_reply("好的") == "收到。"
    assert _resolve_smalltalk_text_reply("ok") == "收到。"


def test_resolve_smalltalk_text_reply_handles_thanks_text():
    assert _resolve_smalltalk_text_reply("谢谢") == "不客气。"
    assert _resolve_smalltalk_text_reply("谢谢啦") == "不客气。"


def test_resolve_smalltalk_text_reply_ignores_real_content():
    assert _resolve_smalltalk_text_reply("这是需要审核的正文第一段内容") is None


def test_review_text_mock_llm(monkeypatch):
    """文字审核 mock LLM 流程."""
    import asyncio

    class _FakeBlock:
        def __init__(self, text: str):
            self.text = text

    class _FakeMessage:
        def __init__(self, text: str):
            self.content = [_FakeBlock(text)]

    class _FakeMessages:
        def __init__(self, text: str):
            self._text = text

        def create(self, **_: object) -> _FakeMessage:
            return _FakeMessage(self._text)

    class _FakeClient:
        def __init__(self, text: str):
            self.messages = _FakeMessages(text)

    monkeypatch.setattr(
        "app.review.general_reviewer.build_anthropic_client",
        lambda: (_FakeClient('{"issues": []}'), "fake-model"),
    )

    config = ReviewConfig(
        wecom_bot_id="x",
        wecom_bot_secret="x",
        rules_path=Path("app/data/rules.md"),
        reviews_dir=Path("data/reviews"),
        logs_dir=Path("data/logs"),
        admin_user_id="",
        admin_name="",
        notification_cooldown=300,
        direct_admin_notifications=False,
        require_registration=False,
    )

    result = asyncio.run(_review_text("第一段。\n\n第二段。", config))
    assert "文字消息" in result
    assert "未发现低级错误" in result


def test_build_user_review_reply_returns_none_when_findings_exist():
    result = ReviewResult(
        findings=[
            Finding(
                rule_id="halfmonthly-leader-title",
                paragraph_index=1,
                line_number=2,
                original_text="黄黎明,行长",
                description="当前信息已采用党内职务口径，黄黎明应补充'党委副书记'",
                target_text="黄黎明",
            )
        ],
        total_rules=12,
        passed_rules=11,
        filename="半月报.docx",
    )

    reply = build_user_review_reply(result, "半月报.docx", doc_type=DocumentType.HALF_MONTHLY)

    assert reply is None


def test_build_user_review_reply_returns_pass_message_when_no_findings():
    result = ReviewResult(
        findings=[],
        total_rules=12,
        passed_rules=12,
        filename="内参周报.docx",
    )

    reply = build_user_review_reply(result, "内参周报.docx", doc_type=DocumentType.NEI_CAN)

    assert reply is not None
    assert "没有发现问题" in reply


def test_run_neican_review_starts_phase2_before_phase1_finishes():
    """内参审核应在等待第一阶段结果时提前启动第二阶段。"""
    import asyncio

    phase2_started = asyncio.Event()

    async def fake_phase1(paragraphs, rules_text, filename, file_path=None):
        await asyncio.wait_for(phase2_started.wait(), timeout=0.1)
        return ReviewResult(
            findings=[],
            total_rules=10,
            passed_rules=10,
            filename=filename,
        )

    async def fake_phase2(paragraphs, rules_text, filename):
        phase2_started.set()
        await asyncio.sleep(0)
        return ReviewResult(
            findings=[],
            total_rules=5,
            passed_rules=5,
            filename=filename,
        )

    async def _run():
        phase1_result, phase2_task, timings = await _start_neican_review(
            phase1_runner=fake_phase1,
            phase2_runner=fake_phase2,
            paragraphs=["标题", "正文"],
            rules_text="",
            filename="demo.docx",
            file_path=None,
        )
        phase2_result, phase2_ms = await phase2_task
        return phase1_result, phase2_result, timings, phase2_ms

    phase1_result, phase2_result, timings, phase2_ms = asyncio.run(_run())

    assert phase1_result.total_rules == 10
    assert phase2_result.total_rules == 5
    assert timings["phase1_ms"] >= 0
    assert timings["wall_start"] > 0
    assert phase2_ms >= 0


def test_run_neican_review_cancels_phase2_when_phase1_fails():
    """第一阶段失败时，应取消后台第二阶段任务，避免继续占用模型调用。"""
    import asyncio

    phase2_cancelled = asyncio.Event()

    async def fake_phase1(paragraphs, rules_text, filename, file_path=None):
        raise RuntimeError("phase1 boom")

    async def fake_phase2(paragraphs, rules_text, filename):
        try:
            await asyncio.sleep(10)
        except asyncio.CancelledError:
            phase2_cancelled.set()
            raise

    try:
        asyncio.run(
            _start_neican_review(
                phase1_runner=fake_phase1,
                phase2_runner=fake_phase2,
                paragraphs=["标题", "正文"],
                rules_text="",
                filename="demo.docx",
                file_path=None,
            )
        )
        assert False, "应抛出第一阶段异常"
    except RuntimeError as exc:
        assert str(exc) == "phase1 boom"

    assert phase2_cancelled.is_set()


def test_processing_failure_reply_hides_internal_exception_details():
    reply = _build_processing_failure_user_reply(
        "文件解析",
        RuntimeError("failed at /private/tmp/job-001 req_id=secret-request"),
    )

    assert "文件解析失败" in reply
    assert "已经提醒管理员" in reply
    assert "/private/tmp" not in reply
    assert "req_id" not in reply


# ============================================================
# 测试入口
# ============================================================

if __name__ == "__main__":
    print("=" * 60)
    print("智能审核 Bot - 存档与配置测试")
    print("=" * 60)
    print()

    tests = [
        test_is_docx_filename_accept,
        test_is_docx_filename_reject,
        test_save_review_creates_directory,
        test_save_review_increments_index,
        test_save_review_sanitizes_filename,
        test_load_config_from_env_file,
        test_load_config_missing_required,
        test_split_text_into_paragraphs,
        test_extract_text_content,
    ]

    passed = 0
    failed = 0
    for t in tests:
        try:
            t()
            passed += 1
        except AssertionError as e:
            print(f"❌ {t.__name__}: FAIL")
            print(f"   {e}")
            failed += 1
        except Exception as e:
            print(f"❌ {t.__name__}: ERROR")
            print(f"   {type(e).__name__}: {e}")
            failed += 1

    print()
    print("=" * 60)
    print(f"测试结果: {passed} 通过, {failed} 失败,共 {len(tests)} 个")
    print("=" * 60)

    sys.exit(0 if failed == 0 else 1)
