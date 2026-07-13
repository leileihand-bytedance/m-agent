"""智能审核模块 - 端到端 + LLM 输出解析测试.

测试范围(LLM 模式下):
  1. LLM 输出 JSON 解析(各种边界情况)
  2. formatter(纯文本输出)
  3. 端到端(需要 LLM,可能慢/可能因模型不同结果不同)

注意:不再测单条规则(LLM 自己判断,代码里没有规则实现)。
"""

from __future__ import annotations

import json
import os
import sys
import tempfile
import zipfile
from pathlib import Path

_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_ROOT))

from app.review import load_rules, review_text, parse_docx  # noqa: E402
from app.review.reviewer import (  # noqa: E402
    Finding,
    ReviewResult,
    _build_review_document,
    _check_weekly_body_format,
    _parse_llm_output,
    check_section_mismatch,
)
from app.review.output_formatter import format_review_result  # noqa: E402


# 启动时从 .env 加载环境变量(端到端测试需要)
def _load_env():
    env_path = _ROOT / ".env"
    if not env_path.exists():
        return
    for line in env_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        k, _, v = line.partition("=")
        k, v = k.strip(), v.strip()
        # 如果环境变量里没设,补上
        if k and k not in os.environ:
            os.environ[k] = v


_load_env()

def _make_synthetic_weekly_format_doc(path: Path) -> None:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.first_line_indent = Cm(0.85)

    for text in ("内部资料", "示例周报", "第1期", "主编：测试用户"):
        doc.add_paragraph(text)

    section = doc.add_paragraph("党政要闻")
    section.paragraph_format.line_spacing = 1.15
    section_run = section.runs[0]
    section_run.font.name = "黑体"
    section_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")
    section_run.font.size = Pt(18)

    title = doc.add_paragraph("示例政策会议召开", style="Heading 3")
    title_run = title.runs[0]
    title_run.font.name = "黑体"
    title_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")
    doc.add_paragraph("会议部署年度重点工作，并提出后续安排。")
    doc.save(path)


# ============================================================
# 测试 1: LLM 输出解析
# ============================================================

def test_parse_llm_output_pure_json():
    """纯 JSON 输入(语义规则)."""
    output = '{"issues": [{"paragraph_index": 0, "rule_id": "title-truncated", "original_text": "习近平对朝鲜进", "description": "标题被截断"}]}'
    paragraphs = ["习近平对朝鲜进", "正常段"]
    findings, reasoning = _parse_llm_output(output, paragraphs)
    assert len(findings) == 1
    assert findings[0].rule_id == "title-truncated"
    assert findings[0].paragraph_index == 0
    assert findings[0].original_text == "习近平对朝鲜进"
    print("✅ test_parse_llm_output_pure_json: 纯 JSON 解析正确")


def test_parse_llm_output_empty():
    """空 issues."""
    output = '{"issues": []}'
    paragraphs = ["任何内容"]
    findings, reasoning = _parse_llm_output(output, paragraphs)
    assert findings == []
    print("✅ test_parse_llm_output_empty: 空 issues 正确")


def test_parse_llm_output_with_markdown_wrapper():
    """带 ```json 包裹的 JSON."""
    output = '```json\n{"issues": [{"paragraph_index": 1, "rule_id": "content-mismatch", "original_text": "标题正文", "description": "标题正文不匹配"}]}\n```'
    paragraphs = ["第一段", "标题正文", "第三段"]
    findings, reasoning = _parse_llm_output(output, paragraphs)
    assert len(findings) == 1
    assert findings[0].rule_id == "content-mismatch"
    assert findings[0].paragraph_index == 1
    print("✅ test_parse_llm_output_with_markdown_wrapper: markdown 包裹正确剥离")


def test_parse_llm_output_with_explanation():
    """LLM 在 JSON 前后加了说明文字(常见!),应能提取 JSON 部分."""
    output = """好的,我审核完了。文档中存在以下问题:

{"issues": [{"paragraph_index": 0, "rule_id": "title-truncated", "original_text": "习近平对朝鲜进", "description": "标题被截断"}]}

希望对你有帮助。"""
    paragraphs = ["习近平对朝鲜进"]
    findings, reasoning = _parse_llm_output(output, paragraphs)
    assert len(findings) == 1, f"应能提取 JSON,实际: {findings}"
    print("✅ test_parse_llm_output_with_explanation: 前后说明文字被正确剥离")


def test_parse_llm_output_invalid_json():
    """无效 JSON 应返回空(让上层处理)."""
    output = "这不是 JSON,只是普通文本"
    paragraphs = ["任何内容"]
    findings, reasoning = _parse_llm_output(output, paragraphs)
    assert findings == []
    print("✅ test_parse_llm_output_invalid_json: 无效 JSON 返回空")


def test_parse_llm_output_index_out_of_range():
    """paragraph_index 越界应被过滤."""
    output = '{"issues": [{"paragraph_index": 99, "rule_id": "title-truncated", "original_text": "x", "description": "y"}]}'
    paragraphs = ["只有一段"]
    findings, reasoning = _parse_llm_output(output, paragraphs)
    assert findings == [], f"越界 index 应被过滤,实际: {findings}"
    print("✅ test_parse_llm_output_index_out_of_range: 越界 index 被过滤")


def test_parse_llm_output_negative_index():
    """paragraph_index 负数应被过滤."""
    output = '{"issues": [{"paragraph_index": -1, "rule_id": "title-truncated", "original_text": "y", "description": "z"}]}'
    paragraphs = ["段落1"]
    findings, reasoning = _parse_llm_output(output, paragraphs)
    assert findings == []
    print("✅ test_parse_llm_output_negative_index: 负数 index 被过滤")


def test_parse_llm_output_multiple_issues():
    """多条问题(语义规则)."""
    output = json.dumps({
        "issues": [
            {"paragraph_index": 0, "rule_id": "title-truncated", "original_text": "习近平对朝鲜进", "description": "标题被截断"},
            {"paragraph_index": 2, "rule_id": "content-mismatch", "original_text": "标题说A正文说B", "description": "内容不匹配"},
            {"paragraph_index": 5, "rule_id": "content-incomplete", "original_text": "正文截断", "description": "语义不完整"},
        ]
    })
    paragraphs = ["习近平对朝鲜进", "正常", "标题说A正文说B", "正常", "正常", "正文截断"]
    findings, reasoning = _parse_llm_output(output, paragraphs)
    assert len(findings) == 3
    assert {f.rule_id for f in findings} == {"title-truncated", "content-mismatch", "content-incomplete"}
    print(f"✅ test_parse_llm_output_multiple_issues: 3 条问题正确解析")


# ============================================================
# 测试 2: formatter(纯文本输出格式)
# ============================================================

def test_formatter_no_issues():
    """没问题."""
    result = ReviewResult(
        findings=[],
        total_rules=13,
        passed_rules=13,
        filename="测试.docx",
    )
    output = format_review_result(result, "测试.docx")
    assert "测试.docx" in output
    assert "未发现低级错误" in output
    print("✅ test_formatter_no_issues: 无问题格式正确")


def test_formatter_with_issues():
    """有问题."""
    result = ReviewResult(
        findings=[
            Finding(
                rule_id="title-truncated",
                paragraph_index=0,
                line_number=1,
                original_text="习近平对朝鲜进行国事访",
                description="标题被截断,缺'问'",
            ),
            Finding(
                rule_id="content-mismatch",
                paragraph_index=1,
                line_number=2,
                original_text="国务院政策例行吹风会介绍推进城市更新工作有关情况",
                description="标题说政策吹风会,正文说求是杂志文章,完全不同",
            ),
        ],
        total_rules=13,
        passed_rules=11,
        filename="报告.docx",
    )
    output = format_review_result(result, "报告.docx")
    assert "📄《报告.docx》" in output
    assert "错误1" in output
    assert "标题截断" in output
    assert "标题正文不匹配" in output
    # 新要求:不说"第 X 行"了(冗余)
    assert "第1行" not in output, f"不应该再出现'第 X 行'信息: {output}"
    assert "第2行" not in output, f"不应该再出现'第 X 行'信息: {output}"
    print("✅ test_formatter_with_issues: 多问题格式正确,无冗余'第 X 行'")


def test_formatter_shows_problem_clearly():
    """formatter 应该直接说问题是什么,不是只给位置."""
    result = ReviewResult(
        findings=[
            Finding(
                rule_id="content-mismatch",
                paragraph_index=37,
                line_number=37,
                original_text="丁向群同志任国家金融监督管理总局党委书记",
                description="标题主题与正文内容不匹配:标题说'丁向群任党委书记',正文却在讲央行档案管理,两完全不相关",
            ),
        ],
        total_rules=12,
        passed_rules=11,
        filename="丁向群.docx",
    )
    output = format_review_result(result, "丁向群.docx")
    # 必须包含:段号 + 规则 + 问题描述
    assert "【标题正文不匹配】" in output
    assert "标题正文不匹配" in output
    assert "丁向群" in output
    assert "不匹配" in output or "不相关" in output
    # 原标题片段要展示
    assert "丁向群同志任" in output
    print("✅ test_formatter_shows_problem_clearly: formatter 直接说明问题")


def test_formatter_llm_error():
    """LLM 调用失败时 formatter 怎么显示."""
    result = ReviewResult(
        findings=[Finding(
            rule_id="__llm_error__",
            paragraph_index=0,
            line_number=1,
            original_text="(LLM 调用失败)",
            description="LLM 调用失败:timeout",
        )],
        total_rules=13,
        passed_rules=0,
        filename="x.docx",
    )
    output = format_review_result(result, "x.docx")
    assert "LLM 调用失败" in output
    print("✅ test_formatter_llm_error: LLM 错误信息正确显示")


def test_formatter_more_than_10_issues():
    """超过 10 条问题:只显示前 10 条 + 提示还有更多."""
    findings = [
        Finding(
            rule_id=f"rule-{i}",
            paragraph_index=0,
            line_number=1,
            original_text=f"段{i}",
            description=f"问题{i}",
        )
        for i in range(15)
    ]
    result = ReviewResult(
        findings=findings,
        total_rules=13,
        passed_rules=0,
        filename="many.docx",
    )
    output = format_review_result(result, "many.docx", max_findings=10)
    assert "错误10" in output, f"应显示到第10条: {output}"
    # 前 10 条应被显示(问题 0~9)
    for i in range(10):
        assert f"问题{i}" in output, f"问题{i} 应显示在输出中"
    print("✅ test_formatter_more_than_10_issues: 15 条问题只显示前 10 条")


# ============================================================
# 测试 3: rules.md 加载
# ============================================================

def test_load_rules_returns_text():
    """load_rules 现在返回字符串(整篇 rules.md 文本)."""
    text = load_rules(Path("app/data/rules.md"))
    assert isinstance(text, str)
    assert "智能审核规则库" in text
    assert "title-truncated" in text
    print(f"✅ test_load_rules_returns_text: 加载了 {len(text)} 字规则文本")


def test_weekly_body_format_ignores_cover_paragraphs_when_doc_has_no_toc():
    with tempfile.TemporaryDirectory() as tmpdir:
        path = Path(tmpdir) / "示例周报.docx"
        _make_synthetic_weekly_format_doc(path)
        parsed = parse_docx(path)

        findings = _check_weekly_body_format(parsed.paragraphs, path)

    cover_findings = [
        f for f in findings
        if f.paragraph_index in {0, 1, 2, 3}
    ]
    assert not cover_findings, (
        "无目录的内参周报，格式审核不应把封面/期号/主编信息当成正文范围。\n"
        f"实际命中: {[(f.paragraph_index, f.description) for f in cover_findings]}"
    )


def test_weekly_body_format_accepts_template_body_indent():
    with tempfile.TemporaryDirectory() as tmpdir:
        path = Path(tmpdir) / "示例周报.docx"
        _make_synthetic_weekly_format_doc(path)
        parsed = parse_docx(path)

        findings = _check_weekly_body_format(parsed.paragraphs, path)

    indent_findings = [
        f for f in findings
        if "正文首行缩进" in f.description
    ]
    assert not indent_findings, (
        "合成模板里的正文首行缩进应视为正确格式，不应整批误报。\n"
        f"实际命中: {[(f.paragraph_index, f.description) for f in indent_findings[:5]]}"
    )


def test_build_review_document_groups_prompt_line_into_previous_entry():
    paragraphs = [
        "目录",
        "市场观察2",
        "前沿观点3",
        "市场观察",
        "示例市场标题",
        "示例市场正文。",
        "前沿观点",
        "示例观点标题",
        "观点背景说明。",
        "原文：以下为补充材料",
        "补充材料正文。",
    ]

    document = _build_review_document(paragraphs)

    titles = [entry.title_text for entry in document.entries]
    assert "原文：以下为补充材料" not in titles, (
        "“原文：”提示行应并入上一条正文，不应被当成独立新闻标题。"
    )

    entry = next(item for item in document.entries if item.title_text == "示例观点标题")
    assert "原文：以下为补充材料" in entry.body_paragraphs


def test_check_section_mismatch_accepts_tongye_dongtai_variant():
    paragraphs = ["同业动态", "某商业银行推出示例产品", "该产品用于验证板块识别。"]

    findings = check_section_mismatch(paragraphs)

    wrong_bank_findings = [
        f for f in findings
        if f.original_text.startswith("某商业银行推出示例产品")
    ]
    assert not wrong_bank_findings, (
        "“同业动态”应视为“同业动向”的合法板块名，"
        "不应把民营银行条目误报为放错板块。\n"
        f"实际命中: {[(f.paragraph_index, f.description) for f in wrong_bank_findings]}"
    )


# ============================================================
# 测试 4: 端到端(构造一个明显有错的 .docx,跑真 LLM)
# 注:这个测试需要 LLM API,可能慢或被 rate-limit
# ============================================================

def _make_fake_docx(path, paragraphs: list[str]):
    """构造一个最小的 .docx."""
    paras_xml = "".join(
        f'<w:p><w:r><w:t>{p}</w:t></w:r></w:p>'
        for p in paragraphs
    )
    document_xml = f'''<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>{paras_xml}</w:body>
</w:document>'''
    content_types_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>'''
    rels_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>'''
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types_xml)
        z.writestr("_rels/.rels", rels_xml)
        z.writestr("word/document.xml", document_xml)


def test_end_to_end_with_clear_errors():
    """端到端:构造明显有错的文档,LLM 应该至少能识别一部分错误.

    注:LLM 偶尔不稳定(可能空 issues),所以这里用重试机制(最多 3 次)。
    """
    import time
    test_path = "/tmp/test_llm_e2e.docx"
    _make_fake_docx(test_path, [
        "我们要的的加快推进工作,实现既定目标。",
        "今天,我去调研企业。",
    ])
    parsed = parse_docx(Path(test_path))
    rules_text = load_rules(Path("app/data/rules.md"))

    # 重试 2 次(LLM 偶发返回空,降低测试超时风险)
    hit_rule_ids = set()
    for attempt in range(2):
        result = review_text(parsed.paragraphs, rules_text, "test.docx")
        hit_rule_ids = {f.rule_id for f in result.findings if not f.rule_id.startswith("__")}
        if hit_rule_ids:
            break
        time.sleep(1)

    assert hit_rule_ids, (
        f"LLM 至少应该识别一条错误,3 次重试都返回空。"
        f"\n所有发现:{[(f.rule_id, f.description) for f in result.findings]}"
    )
    print(f"✅ test_end_to_end_with_clear_errors: LLM 识别出 {hit_rule_ids} (重试 {attempt+1} 次)")


def test_end_to_end_with_no_errors():
    """端到端:构造正常文档,LLM 应该输出 0 issues."""
    test_path = "/tmp/test_llm_no_errors.docx"
    _make_fake_docx(test_path, [
        "今天我们召开了专题会议,研究部署下一阶段工作。",
        "各部门负责人结合实际情况,提出了具体的落实措施。",
    ])
    parsed = parse_docx(Path(test_path))
    rules_text = load_rules(Path("app/data/rules.md"))
    result = review_text(parsed.paragraphs, rules_text, "clean.docx")

    # 注意:LLM 可能还会报"今天,我去"那种小问题,这里只验证不会爆 LLM 错误
    llm_errors = [f for f in result.findings if f.rule_id.startswith("__")]
    assert not llm_errors, f"LLM 调用不应出错: {[f.description for f in llm_errors]}"
    print(f"✅ test_end_to_end_with_no_errors: LLM 正常返回 {len(result.findings)} 条(可能为 0)")


def test_end_to_end_content_mismatch():
    """端到端：合成标题讲人事、正文讲档案管理的错配样例。"""
    test_path = "/tmp/test_llm_content_mismatch.docx"
    _make_fake_docx(test_path, [
        # 标题段(短)
        "示例机构任命新负责人",
        # 正文段
        "示例部门发布档案管理办法，正文介绍归档范围、保管期限、信息化建设和安全管理要求，与负责人任命无关。",
    ])
    parsed = parse_docx(Path(test_path))
    rules_text = load_rules(Path("app/data/rules.md"))
    result = review_text(parsed.paragraphs, rules_text, "synthetic_mismatch.docx")

    # LLM 应该识别 content-mismatch
    hit_rule_ids = {f.rule_id for f in result.findings if not f.rule_id.startswith("__")}
    assert "content-mismatch" in hit_rule_ids, (
        f"LLM 应该识别内容不匹配,实际命中:{hit_rule_ids}\n"
        f"所有发现:{[(f.rule_id, f.description) for f in result.findings]}"
    )
    print(f"✅ test_end_to_end_content_mismatch: LLM 识别出 {hit_rule_ids}")


def test_end_to_end_content_match():
    """反例:标题和正文一致,LLM 不应报 content-mismatch."""
    test_path = "/tmp/test_llm_content_match.docx"
    _make_fake_docx(test_path, [
        "示例部门发布档案管理办法",
        "示例部门发布档案管理办法，正文介绍归档范围、保管期限和安全管理要求。",
    ])
    parsed = parse_docx(Path(test_path))
    rules_text = load_rules(Path("app/data/rules.md"))
    result = review_text(parsed.paragraphs, rules_text, "match.docx")

    # 不应该报 content-mismatch(标题正文一致)
    mismatch = [f for f in result.findings if f.rule_id == "content-mismatch"]
    assert not mismatch, (
        f"标题正文一致不应报 content-mismatch,实际报了:{mismatch}"
    )
    print(f"✅ test_end_to_end_content_match: 标题正文一致,未误报")


def test_end_to_end_multiple_content_mismatches():
    """端到端:多组标题-正文错配,LLM 应识别多条 content-mismatch.

    这是用于验证多组标题和正文被混搭的合成场景。
    """
    import time
    test_path = "/tmp/test_llm_multiple_mismatch.docx"
    _make_fake_docx(test_path, [
        # 段 0: 标题 A
        "示例机构任命新负责人",
        # 段 1: 正文却是 B(错配 1)
        "示例部门发布档案管理办法",
        # 段 2: 标题 C
        "示例代表团举行合作会谈",
        # 段 3: 正文却是 D(错配 2)
        "上周示例指数小幅上涨，市场交易保持平稳。",
        # 段 4: 标题 E
        "示例机构负责人会见合作方代表",
        # 段 5: 正文讲的是会见议题(匹配,不报)
        "示例机构负责人会见合作方代表，双方就后续合作安排交换意见。",
    ])
    parsed = parse_docx(Path(test_path))
    rules_text = load_rules(Path("app/data/rules.md"))

    # 重试 2 次
    mismatches = []
    for attempt in range(2):
        result = review_text(parsed.paragraphs, rules_text, "multi_mismatch.docx")
        mismatches = [f for f in result.findings if f.rule_id == "content-mismatch"]
        if len(mismatches) >= 2:
            break
        time.sleep(1)

    assert len(mismatches) >= 2, (
        f"期望 LLM 至少识别 2 处内容错配(段 0、段 2),2 次重试后只报 {len(mismatches)} 处: "
        f"{[f.description[:60] for f in mismatches]}"
    )
    print(f"✅ test_end_to_end_multiple_content_mismatches: LLM 识别出 {len(mismatches)} 处错配 (重试 {attempt+1} 次)")


def test_end_to_end_no_false_positive_on_toc():
    """端到端:目录项不应被误报为 content-mismatch 或 quote-pair.

    这是真实文档里最常见的误报:目录里的短标题末尾带页码数字,LLM 误以为它
    是"应该有书名号"的标题。
    """
    import time
    test_path = "/tmp/test_llm_toc.docx"
    _make_fake_docx(test_path, [
        # 段 0: 文档标题
        "内部资料",
        # 段 1: 周报名
        "微众银行信息内参周报",
        # 段 2: 目录标题
        "目录",
        # 段 3-7: 目录项(短段 + 末尾是数字)
        "《求是》杂志发表习近平总书记重要文章《前瞻布局和发展未来产业》2",
        "习近平同塞尔维亚总统武契奇会谈2",
        "国务院印发《城市更新十五五规划》2",
        "丁向群同志任国家金融监督管理总局党委书记2",
        "何立峰会见德国联邦经济和能源部部长赖歇一行3",
        # 段 8: 章节分类
        "党政要闻",
        # 段 9: 真新闻标题(末尾不带数字)
        "丁向群同志任国家金融监督管理总局党委书记",
        # 段 10: 对应正文(正确匹配)
        "近日，丁向群同志任国家金融监督管理总局党委书记，会议指出，要全面加强党的建设，推动事业高质量发展。",
    ])
    parsed = parse_docx(Path(test_path))
    rules_text = load_rules(Path("app/data/rules.md"))

    # 重试 2 次
    findings = []
    for attempt in range(2):
        result = review_text(parsed.paragraphs, rules_text, "toc.docx")
        findings = result.findings
        # 如果能稳定下来(没有 quote-pair/缺少书名号 类误报),跳出
        if not any("缺少书名号" in f.description for f in findings):
            break
        time.sleep(1)

    # 不应有"缺少书名号"这类误报(目录项本来就不需要书名号)
    book_quote_false_positives = [
        f for f in findings
        if "缺少书名号" in f.description
    ]
    assert not book_quote_false_positives, (
        f"目录项不应被误报为'缺少书名号',实际报了 {len(book_quote_false_positives)} 条: "
        f"{[f.description[:60] for f in book_quote_false_positives]}"
    )

    # 段 8 (党政要闻) 是章节分类,不应被报 content-mismatch
    mismatch_in_toc = [
        f for f in findings
        if f.rule_id == "content-mismatch" and f.paragraph_index in (0, 1, 2, 3, 4, 5, 6, 7, 8)
    ]
    assert not mismatch_in_toc, (
        f"目录项/章节分类不应报 content-mismatch,实际: "
        f"{[(f.paragraph_index, f.description[:60]) for f in mismatch_in_toc]}"
    )

    print(f"✅ test_end_to_end_no_false_positive_on_toc: 目录项无误报 ({len(findings)} 条发现均为真问题)")


def test_end_to_end_synthetic_mismatches():
    """端到端：合成标题正文错配和正文截断样例。"""
    import time
    test_path = "/tmp/test_llm_synthetic_mismatches.docx"
    _make_fake_docx(test_path, [
        # 段 0: 真新闻标题 1
        "示例项目召开年度推进会",
        # 段 1: 正文却是产品发布说明(错配 1)
        "示例机构发布一项测试产品，正文主要介绍产品功能、使用对象和操作流程，与年度推进会无关。",
        # 段 2: 真新闻标题 2
        "示例机构部署下一阶段工作",
        # 段 3: 正文戛然而止(截断 1)
        "会议部署了流程优化、协同管理和服务提升等任务，并强调下一阶段将。",
        # 段 4: 一个完整句子的对照(不应报截断)
        "市场观察方面，本周示例指数小幅波动，参与者保持谨慎。",
    ])
    parsed = parse_docx(Path(test_path))
    rules_text = load_rules(Path("app/data/rules.md"))

    # 重试 2 次
    findings = []
    for attempt in range(2):
        result = review_text(parsed.paragraphs, rules_text, "synthetic_mismatches.docx")
        findings = result.findings
        content_mismatch = [f for f in findings if f.rule_id == "content-mismatch"]
        content_incomplete = [f for f in findings if f.rule_id == "content-incomplete"]
        # 都命中即可
        if content_mismatch and content_incomplete:
            break
        time.sleep(1)

    # 验证 1: 段 0 标题应该被报 content-mismatch
    mismatch_on_title = [
        f for f in findings
        if f.rule_id == "content-mismatch" and f.paragraph_index == 0
    ]
    assert mismatch_on_title, (
        f"段 0 标题应报 content-mismatch(标题讲推进会,正文讲产品),实际没报。\n"
        f"所有发现:{[(f.paragraph_index, f.rule_id, f.description[:60]) for f in findings]}"
    )

    # 验证 2: 段 3 正文应该被报 content-incomplete(戛然而止)
    incomplete_on_truncated = [
        f for f in findings
        if f.rule_id == "content-incomplete" and f.paragraph_index == 3
    ]
    assert incomplete_on_truncated, (
        f"段 3 正文应报 content-incomplete(戛然而止),实际没报。\n"
        f"所有发现:{[(f.paragraph_index, f.rule_id, f.description[:60]) for f in findings]}"
    )

    # 验证 3: 段 4 完整句不应被报 content-incomplete
    incomplete_on_complete = [
        f for f in findings
        if f.rule_id == "content-incomplete" and f.paragraph_index == 4
    ]
    assert not incomplete_on_complete, (
        f"段 4 是完整句,不应报 content-incomplete,实际报了:\n"
        f"{[f.description[:80] for f in incomplete_on_complete]}"
    )

    print(f"✅ test_end_to_end_synthetic_mismatches: LLM 识别 {len(mismatch_on_title)} 错配 + {len(incomplete_on_truncated)} 截断 (重试 {attempt+1} 次)")


def test_end_to_end_title_truncated():
    """端到端：合成标题截断样例。

    验证 LLM 能识别 3 种标题截断场景:
    1. 后面紧跟完全无关的段
    2. 后面是长段正文
    3. 单段标题,后面是页脚

    注:测试用例不要把"截断版"和"完整版"挨着放,LLM 会困惑。
    """
    import time
    test_path = "/tmp/test_llm_title_truncated.docx"
    _make_fake_docx(test_path, [
        # 段 0: 完整标题 + 对应长正文(对照:不应报截断)
        "示例机构与合作方举行年度会谈",
        "双方围绕年度合作计划、项目推进和后续安排深入交换意见，并明确了下一阶段工作重点。",
        # 段 2: 截断标题 + 紧跟长正文(真截断:应报)
        "示例机构与合作方举",
        "双方围绕年度合作计划、项目推进和后续安排深入交换意见，并明确了下一阶段工作重点。",
        # 段 4: 完整标题 + 对应正文(对照:不应报)
        "四月示例市场规模明显回升",
        "数据显示，截至四月末，示例市场规模较上月有所增加。",
        # 段 6: 截断标题 + 紧跟正文(真截断:应报)
        "四月示例市场规模明显回",
        "数据显示，截至四月末，示例市场规模较上月有所增加。",
        # 段 8: 章节分类(对照:不报)
        "市场观察",
        # 段 9: 截断标题 + 紧跟另一章节(真截断:应报)
        "示例市场波动加剧，相关指",
        "市场观察方面，本周示例指数小幅波动。",
    ])
    parsed = parse_docx(Path(test_path))
    rules_text = load_rules(Path("app/data/rules.md"))

    # 重试 2 次
    findings = []
    for attempt in range(2):
        result = review_text(parsed.paragraphs, rules_text, "title_truncated.docx")
        findings = result.findings
        title_truncated = [
            f for f in findings
            if f.rule_id == "title-truncated"
        ]
        if len(title_truncated) >= 3:
            break
        time.sleep(1)

    # 段 2、6、9 应被报 title-truncated
    truncated_indices = {
        f.paragraph_index for f in findings
        if f.rule_id == "title-truncated"
    }
    expected_truncated = {2, 6, 9}
    missing = expected_truncated - truncated_indices
    assert not missing, (
        f"应报 title-truncated 的段: {expected_truncated},实际报: {truncated_indices},缺: {missing}"
    )

    # 段 0、4 完整标题不应误报
    false_positives = {
        f.paragraph_index for f in findings
        if f.rule_id == "title-truncated" and f.paragraph_index in (0, 4)
    }
    assert not false_positives, (
        f"完整标题不应报 title-truncated,实际误报: {false_positives}"
    )

    print(f"✅ test_end_to_end_title_truncated: 3 处标题截断全部识别 (重试 {attempt+1} 次)")


def test_end_to_end_toc_with_number():
    """端到端：合成目录项不应带“一、二、三”序号的样例。

    示例周报目录项出现:
      - "一、党政要闻2"
      - "三、市场观察3"(跳号)
      - "四、前沿观点4"
    应该被报。
    """
    import time
    test_path = "/tmp/test_llm_toc_with_number.docx"
    _make_fake_docx(test_path, [
        # 段 0: 内部资料
        "内部资料",
        # 段 1: 周报名
        "微众银行信息内参周报",
        # 段 2: 目录
        "目录",
        # 段 3: 目录项带"一、"
        "一、党政要闻2",
        # 段 4: 目录项不带序号(对照)
        "监管动态3",
        # 段 5: 目录项带"三、"(跳号)
        "三、市场观察3",
        # 段 6: 目录项带"四、"
        "四、前沿观点4",
        # 段 7: 章节分类(对照)
        "党政要闻",
        # 段 8: 正文
        "本周重要新闻:6月1日,求是杂志发表重要文章。",
    ])
    parsed = parse_docx(Path(test_path))
    rules_text = load_rules(Path("app/data/rules.md"))

    # 重试 2 次
    findings = []
    for attempt in range(2):
        result = review_text(parsed.paragraphs, rules_text, "toc_with_number.docx")
        findings = result.findings
        toc_problems = [
            f for f in findings
            if "目录" in f.description or "toc" in f.rule_id.lower()
        ]
        if len(toc_problems) >= 3:
            break
        time.sleep(1)

    # 段 3、5、6 应该被报(目录项带序号)
    toc_indices = {
        f.paragraph_index for f in findings
        if f.paragraph_index in (3, 5, 6) and (
            "目录" in f.description or "序号" in f.description
            or "一、" in f.description or "三、" in f.description
        )
    }
    assert len(toc_indices) >= 2, (
        f"应报至少 2 处目录项带序号的问题,实际:\n"
        f"{[(f.paragraph_index, f.rule_id, f.description[:80]) for f in findings]}"
    )

    # 段 4(纯章节名"监管动态")不应误报
    false_positives = [
        f for f in findings
        if f.paragraph_index == 4
    ]
    assert not false_positives, (
        f"段 4 是纯章节名,不应误报,实际:\n"
        f"{[(f.rule_id, f.description[:80]) for f in false_positives]}"
    )

    print(f"✅ test_end_to_end_toc_with_number: {len(toc_indices)} 处目录项带序号被识别 (重试 {attempt+1} 次)")


# ============================================================
# 测试 5: phase1/phase2 规则分组
# ============================================================

def test_phase1_rules_only():
    """Phase1 只返回格式规则和基础内容规则."""
    from app.review.reviewer import PHASE1_RULES, PHASE2_RULES
    assert "title-truncated" in PHASE1_RULES
    assert "content-mismatch" in PHASE1_RULES
    assert "content-incomplete" in PHASE1_RULES
    assert "toc-mismatch" not in PHASE1_RULES
    assert "content-out-of-scope" not in PHASE1_RULES


def test_phase2_rules_only():
    """Phase2 只返回深度内容规则."""
    from app.review.reviewer import PHASE1_RULES, PHASE2_RULES
    assert "toc-mismatch" in PHASE2_RULES
    assert "content-out-of-scope" in PHASE2_RULES
    assert "content-wrong-section" in PHASE2_RULES
    assert "content-duplicate" in PHASE2_RULES
    assert "content-outdated" in PHASE2_RULES
    assert "title-truncated" not in PHASE2_RULES


def test_is_news_title_allows_question_style_headline():
    from app.review.reviewer import _is_news_title

    assert _is_news_title('国内距离“零利率”还有多远？')


# ============================================================
# 主入口
# ============================================================

if __name__ == "__main__":
    print("=" * 60)
    print("智能审核模块 - LLM 模式测试")
    print("=" * 60)
    print()

    # 必须测试
    must_tests = [
        test_parse_llm_output_pure_json,
        test_parse_llm_output_empty,
        test_parse_llm_output_with_markdown_wrapper,
        test_parse_llm_output_with_explanation,
        test_parse_llm_output_invalid_json,
        test_parse_llm_output_index_out_of_range,
        test_parse_llm_output_negative_index,
        test_parse_llm_output_multiple_issues,
        test_formatter_no_issues,
        test_formatter_with_issues,
        test_formatter_shows_problem_clearly,
        test_formatter_llm_error,
        test_formatter_more_than_10_issues,
        test_load_rules_returns_text,
        test_weekly_body_format_ignores_cover_paragraphs_when_doc_has_no_toc,
        test_weekly_body_format_accepts_template_body_indent,
        test_is_news_title_allows_question_style_headline,
    ]

    # 端到端测试(可能需要 LLM)
    e2e_tests = [
        test_end_to_end_with_clear_errors,
        test_end_to_end_with_no_errors,
        test_end_to_end_content_mismatch,
        test_end_to_end_content_match,
        test_end_to_end_multiple_content_mismatches,
        test_end_to_end_no_false_positive_on_toc,
        test_end_to_end_synthetic_mismatches,
        test_end_to_end_title_truncated,
        test_end_to_end_toc_with_number,
    ]

    passed = 0
    failed = 0

    print("--- 必须测试 ---")
    for t in must_tests:
        try:
            t()
            passed += 1
        except AssertionError as e:
            print(f"❌ {t.__name__}: FAIL")
            print(f"   {e}")
            failed += 1
        except Exception as e:
            print(f"❌ {t.__name__}: ERROR")
            print(f"   {type(e).__name__}: {e}")
            failed += 1

    print()
    print("--- 端到端测试(需要 LLM) ---")
    import os
    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("⚠️  ANTHROPIC_API_KEY 未设置,跳过端到端测试")
    else:
        for t in e2e_tests:
            try:
                t()
                passed += 1
            except AssertionError as e:
                print(f"❌ {t.__name__}: FAIL")
                print(f"   {e}")
                failed += 1
            except Exception as e:
                print(f"❌ {t.__name__}: ERROR")
                print(f"   {type(e).__name__}: {e}")
                failed += 1

    print()
    print("=" * 60)
    print(f"测试结果:{passed} 通过,{failed} 失败,共 {len(must_tests) + len(e2e_tests)} 个")
    print("=" * 60)

    sys.exit(0 if failed == 0 else 1)
