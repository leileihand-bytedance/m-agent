from __future__ import annotations

import asyncio
import csv
import json
import os
import shutil
import subprocess
import sys
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

import pytest
from docx import Document

from app.review.quality_evaluation import (
    ReviewRunMetrics,
    ReviewSampleCandidate,
    SelectedReviewCase,
    discover_general_candidates,
    run_baseline,
    select_baseline_cases,
)
from app.review.reviewer import Finding, ReviewResult
from scripts.review_quality import build_parser


def _write_docx(path: Path, paragraphs: list[str], *, with_table: bool = False) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if with_table:
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "项目"
        table.cell(0, 1).text = "内容"
        table.cell(1, 0).text = "附件1"
        table.cell(1, 1).text = "议案意见反馈表"
    document.save(path)


def _candidate(
    name: str,
    *,
    total_chars: int,
    table_count: int = 0,
    has_attachment_reference: bool = False,
    is_questionnaire: bool = False,
) -> ReviewSampleCandidate:
    return ReviewSampleCandidate(
        source_path=Path(f"/external/{name}.docx"),
        source_filename=f"{name}.docx",
        sha256=name * 8,
        total_chars=total_chars,
        total_paragraphs=10,
        table_count=table_count,
        has_attachment_reference=has_attachment_reference,
        is_questionnaire=is_questionnaire,
        modified_at=1.0,
    )


def test_discover_general_candidates_deduplicates_and_extracts_features(
    tmp_path: Path,
) -> None:
    tasks_root = tmp_path / "tasks" / "review"
    ordinary = tasks_root / "2026" / "07" / "001" / "input" / "工作材料.docx"
    duplicate = tasks_root / "2026" / "07" / "002" / "input" / "工作材料副本.docx"
    questionnaire = (
        tasks_root / "2026" / "07" / "003" / "input" / "行业调研问卷.docx"
    )
    inner_report = (
        tasks_root
        / "2026"
        / "07"
        / "004"
        / "input"
        / "微众银行信息内参周报2026年第1期.docx"
    )

    _write_docx(ordinary, ["工作材料", "这是正文。"])
    duplicate.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(ordinary, duplicate)
    _write_docx(questionnaire, ["调研问卷", "请填写附件1。"], with_table=True)
    _write_docx(inner_report, ["内部资料", "微众银行信息内参周报"])

    candidates = discover_general_candidates(tasks_root)

    assert len(candidates) == 2
    assert len({candidate.sha256 for candidate in candidates}) == 2
    questionnaire_candidate = next(
        candidate for candidate in candidates if candidate.is_questionnaire
    )
    assert questionnaire_candidate.table_count == 1
    assert questionnaire_candidate.has_attachment_reference is True
    assert questionnaire_candidate.total_paragraphs >= 4


def test_discover_general_candidates_drops_near_duplicate_long_versions(
    tmp_path: Path,
) -> None:
    tasks_root = tmp_path / "tasks" / "review"
    older = tasks_root / "001" / "input" / "长材料旧版.docx"
    newer = tasks_root / "002" / "input" / "长材料新版.docx"
    repeated = "这是用于近重复识别的较长正文内容。" * 180
    _write_docx(older, ["长材料", repeated])
    _write_docx(newer, ["长材料", repeated.replace("较长正文", "较长的正文", 1)])
    os.utime(older, (1, 1))
    os.utime(newer, (2, 2))

    candidates = discover_general_candidates(tasks_root)

    assert len(candidates) == 1
    assert candidates[0].source_path == newer


def test_discover_recognizes_unnumbered_attachment_section(tmp_path: Path) -> None:
    tasks_root = tmp_path / "tasks" / "review"
    source = tasks_root / "001" / "input" / "通知.docx"
    _write_docx(source, ["工作通知", "请按要求报送材料。", "附件：业务情况统计表"])

    candidates = discover_general_candidates(tasks_root)

    assert len(candidates) == 1
    assert candidates[0].has_attachment_reference is True


def test_select_baseline_cases_covers_five_primary_categories_without_duplicates() -> None:
    candidates = [
        _candidate("questionnaire", total_chars=3000, table_count=8, is_questionnaire=True),
        _candidate("attachment", total_chars=6000, has_attachment_reference=True),
        _candidate("long", total_chars=30000),
        _candidate("table", total_chars=8000, table_count=5),
        _candidate("ordinary", total_chars=5000),
        _candidate("extra", total_chars=1200),
    ]

    selected = select_baseline_cases(candidates, limit=5)

    assert [case.case_id for case in selected] == [
        "G-001",
        "G-002",
        "G-003",
        "G-004",
        "G-005",
    ]
    assert {case.primary_category for case in selected} == {
        "问卷",
        "附件引用",
        "长文",
        "表格",
        "常规材料",
    }
    assert len({case.candidate.sha256 for case in selected}) == 5


def test_review_run_metrics_counts_calls_across_threads() -> None:
    metrics = ReviewRunMetrics()

    with ThreadPoolExecutor(max_workers=8) as executor:
        list(executor.map(lambda _: metrics.record_model_call(), range(500)))
        list(executor.map(lambda _: metrics.record_model_failure(), range(20)))
        list(
            executor.map(
                lambda index: metrics.record_degraded_stage(f"chunk_{index % 2 + 1}"),
                range(20),
            )
        )

    assert metrics.model_calls == 500
    assert metrics.model_failures == 20
    assert metrics.degraded_stages == ("chunk_1", "chunk_2")


def test_run_baseline_keeps_failed_case_and_writes_user_scoring_csv(
    tmp_path: Path,
) -> None:
    data_root = tmp_path / "M-Agent-Files"
    good_path = data_root / "tasks" / "review" / "001" / "input" / "正常材料.docx"
    bad_path = data_root / "tasks" / "review" / "002" / "input" / "失败材料.docx"
    _write_docx(good_path, ["正常材料", "这里有一个问题。"])
    _write_docx(bad_path, ["失败材料", "这里会触发模型失败。"])

    good = ReviewSampleCandidate(
        source_path=good_path,
        source_filename=good_path.name,
        sha256="a" * 64,
        total_chars=20,
        total_paragraphs=2,
        table_count=0,
        has_attachment_reference=False,
        is_questionnaire=False,
        modified_at=1.0,
    )
    bad = ReviewSampleCandidate(
        source_path=bad_path,
        source_filename=bad_path.name,
        sha256="b" * 64,
        total_chars=20,
        total_paragraphs=2,
        table_count=0,
        has_attachment_reference=False,
        is_questionnaire=False,
        modified_at=1.0,
    )
    cases = [
        SelectedReviewCase("G-001", "常规材料", good),
        SelectedReviewCase("G-002", "补充样本", bad),
    ]

    async def fake_reviewer(
        paragraphs: list[str],
        rules_text: str,
        filename: str,
        *,
        metrics: ReviewRunMetrics | None = None,
    ) -> ReviewResult:
        assert rules_text == "审核规则"
        assert metrics is not None
        metrics.record_model_call()
        if filename == bad_path.name:
            raise RuntimeError("模型连接失败")
        return ReviewResult(
            findings=[
                Finding(
                    rule_id="general-grammar",
                    paragraph_index=1,
                    line_number=2,
                    original_text=paragraphs[1],
                    target_text="一个问题",
                    description="测试问题说明",
                )
            ],
            total_rules=1,
            passed_rules=0,
            filename=filename,
        )

    summary = asyncio.run(
        run_baseline(
            cases,
            run_id="test-baseline",
            data_root=data_root,
            output_root=data_root / "evaluations" / "review",
            rules_text="审核规则",
            reviewer=fake_reviewer,
        )
    )

    assert summary.total_cases == 2
    assert summary.completed_cases == 1
    assert summary.failed_cases == 1
    assert summary.total_model_calls == 2
    run_dir = data_root / "evaluations" / "review" / "test-baseline"
    assert (run_dir / "cases" / "G-001" / "marked.docx").exists()
    failed_result = json.loads(
        (run_dir / "cases" / "G-002" / "result.json").read_text(encoding="utf-8")
    )
    assert failed_result["status"] == "failed"
    assert failed_result["error"] == "模型连接失败"

    with (run_dir / "scoring.csv").open(encoding="utf-8-sig", newline="") as source:
        reader = csv.DictReader(source)
        rows = list(reader)
        headers = reader.fieldnames or []
    assert len(rows) == 1
    assert rows[0]["错误原文"] == "这里有一个问题。"
    assert "段落编号" not in headers
    assert "paragraph_index_internal" not in headers


def test_run_baseline_rejects_output_outside_data_root(tmp_path: Path) -> None:
    with pytest.raises(ValueError, match="评测输出必须位于"):
        asyncio.run(
            run_baseline(
                [],
                run_id="unsafe",
                data_root=tmp_path / "M-Agent-Files",
                output_root=tmp_path / "repo-output",
                rules_text="",
            )
        )


def test_run_baseline_escapes_spreadsheet_formula_values(tmp_path: Path) -> None:
    data_root = tmp_path / "M-Agent-Files"
    source = data_root / "tasks" / "review" / "001" / "input" / "公式材料.docx"
    _write_docx(source, ["公式材料", "=2+2"])
    candidate = ReviewSampleCandidate(
        source,
        source.name,
        "d" * 64,
        7,
        2,
        0,
        False,
        False,
        1.0,
    )

    async def fake_reviewer(
        paragraphs: list[str],
        rules_text: str,
        filename: str,
        *,
        metrics: ReviewRunMetrics | None = None,
    ) -> ReviewResult:
        return ReviewResult(
            [Finding("general-grammar", 1, 2, paragraphs[1], "测试", "=2+2")],
            1,
            0,
            filename,
        )

    summary = asyncio.run(
        run_baseline(
            [SelectedReviewCase("G-001", "常规材料", candidate)],
            run_id="csv-safety",
            data_root=data_root,
            output_root=data_root / "evaluations" / "review",
            rules_text="",
            reviewer=fake_reviewer,
        )
    )

    with (summary.run_dir / "scoring.csv").open(
        encoding="utf-8-sig",
        newline="",
    ) as csv_file:
        row = next(csv.DictReader(csv_file))
    assert row["错误原文"] == "'=2+2"
    assert row["错误片段"] == "'=2+2"


def test_run_baseline_marks_unrecovered_review_stage_as_partial_failure(
    tmp_path: Path,
) -> None:
    data_root = tmp_path / "M-Agent-Files"
    source = data_root / "tasks" / "review" / "001" / "input" / "材料.docx"
    _write_docx(source, ["材料", "正文。"])
    candidate = ReviewSampleCandidate(
        source_path=source,
        source_filename=source.name,
        sha256="c" * 64,
        total_chars=6,
        total_paragraphs=2,
        table_count=0,
        has_attachment_reference=False,
        is_questionnaire=False,
        modified_at=1.0,
    )

    async def degraded_reviewer(
        paragraphs: list[str],
        rules_text: str,
        filename: str,
        *,
        metrics: ReviewRunMetrics | None = None,
    ) -> ReviewResult:
        assert metrics is not None
        metrics.record_model_call()
        metrics.record_model_failure()
        metrics.record_degraded_stage("chunk_2")
        return ReviewResult([], 1, 1, filename)

    summary = asyncio.run(
        run_baseline(
            [SelectedReviewCase("G-001", "常规材料", candidate)],
            run_id="partial-baseline",
            data_root=data_root,
            output_root=data_root / "evaluations" / "review",
            rules_text="",
            reviewer=degraded_reviewer,
        )
    )

    result = json.loads(
        (summary.run_dir / "cases" / "G-001" / "result.json").read_text(
            encoding="utf-8"
        )
    )
    assert summary.completed_cases == 0
    assert summary.failed_cases == 1
    assert result["status"] == "partial_failed"
    assert result["model_failures"] == 1
    assert result["degraded_stages"] == ["chunk_2"]


def test_run_baseline_resume_skips_completed_case_and_reruns_failed_case(
    tmp_path: Path,
) -> None:
    data_root = tmp_path / "M-Agent-Files"
    first_path = data_root / "tasks" / "review" / "001" / "input" / "第一份.docx"
    second_path = data_root / "tasks" / "review" / "002" / "input" / "第二份.docx"
    _write_docx(first_path, ["第一份", "正文一。"])
    _write_docx(second_path, ["第二份", "正文二。"])
    cases = [
        SelectedReviewCase(
            "G-001",
            "常规材料",
            ReviewSampleCandidate(
                first_path,
                first_path.name,
                "1" * 64,
                8,
                2,
                0,
                False,
                False,
                1.0,
            ),
        ),
        SelectedReviewCase(
            "G-002",
            "补充样本",
            ReviewSampleCandidate(
                second_path,
                second_path.name,
                "2" * 64,
                8,
                2,
                0,
                False,
                False,
                1.0,
            ),
        ),
    ]
    first_calls: list[str] = []

    async def first_reviewer(
        paragraphs: list[str],
        rules_text: str,
        filename: str,
        *,
        metrics: ReviewRunMetrics | None = None,
    ) -> ReviewResult:
        first_calls.append(filename)
        if filename == second_path.name:
            raise RuntimeError("暂时失败")
        return ReviewResult(
            [Finding("general-grammar", 1, 2, paragraphs[1], "问题一", "正文一")],
            1,
            0,
            filename,
        )

    asyncio.run(
        run_baseline(
            cases,
            run_id="resume-baseline",
            data_root=data_root,
            output_root=data_root / "evaluations" / "review",
            rules_text="",
            reviewer=first_reviewer,
        )
    )
    completed_result_path = (
        data_root
        / "evaluations"
        / "review"
        / "resume-baseline"
        / "cases"
        / "G-001"
        / "result.json"
    )
    legacy_result = json.loads(completed_result_path.read_text(encoding="utf-8"))
    legacy_result.pop("model_failures")
    legacy_result.pop("degraded_stages")
    completed_result_path.write_text(
        json.dumps(legacy_result, ensure_ascii=False),
        encoding="utf-8",
    )
    resumed_calls: list[str] = []

    async def resumed_reviewer(
        paragraphs: list[str],
        rules_text: str,
        filename: str,
        *,
        metrics: ReviewRunMetrics | None = None,
    ) -> ReviewResult:
        resumed_calls.append(filename)
        return ReviewResult(
            [Finding("general-grammar", 1, 2, paragraphs[1], "问题二", "正文二")],
            1,
            0,
            filename,
        )

    summary = asyncio.run(
        run_baseline(
            cases,
            run_id="resume-baseline",
            data_root=data_root,
            output_root=data_root / "evaluations" / "review",
            rules_text="",
            reviewer=resumed_reviewer,
            resume=True,
        )
    )

    assert first_calls == [first_path.name, second_path.name]
    assert resumed_calls == [second_path.name]
    assert summary.completed_cases == 2
    assert summary.failed_cases == 0
    normalized_result = json.loads(completed_result_path.read_text(encoding="utf-8"))
    assert normalized_result["model_failures"] == 0
    assert normalized_result["degraded_stages"] == []
    with (summary.run_dir / "scoring.csv").open(
        encoding="utf-8-sig",
        newline="",
    ) as source:
        assert len(list(csv.DictReader(source))) == 2


def test_review_quality_cli_defaults_to_five_real_cases() -> None:
    args = build_parser().parse_args(["run", "--run-id", "baseline-v1"])

    assert args.command == "run"
    assert args.limit == 5
    assert args.run_id == "baseline-v1"
    assert args.resume is False


def test_review_quality_script_runs_directly_from_outside_project(tmp_path: Path) -> None:
    script_path = Path(__file__).resolve().parents[1] / "scripts" / "review_quality.py"

    completed = subprocess.run(
        [sys.executable, str(script_path), "--help"],
        cwd=tmp_path,
        capture_output=True,
        text=True,
        check=False,
    )

    assert completed.returncode == 0, completed.stderr
    assert "通用审核真实文件质量评测" in completed.stdout
