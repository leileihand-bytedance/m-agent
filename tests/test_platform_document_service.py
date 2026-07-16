from __future__ import annotations

import base64
from io import BytesIO
import json
from pathlib import Path
import zipfile

import pytest
from docx import Document
from docx.shared import Inches
from pypdf import PdfWriter

from app.platform.documents import (
    DocumentFormat,
    DocumentSecurityError,
    DocumentService,
)
from app.platform.documents.parsers.pptx import _chart_text


def _make_docx(path: Path) -> None:
    document = Document()
    document.add_heading("测试标题", level=1)
    document.add_paragraph("第一段正文。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "100亿元"
    section = document.sections[0]
    section.header.paragraphs[0].text = "页眉文字"
    document.save(path)


def _tiny_png() -> BytesIO:
    payload = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
    )
    return BytesIO(payload)


def _make_blank_pdf(path: Path, pages: int = 2) -> None:
    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=595, height=842)
    with path.open("wb") as stream:
        writer.write(stream)


def _make_pptx(path: Path) -> None:
    pptx = pytest.importorskip("pptx")
    presentation = pptx.Presentation()
    slide = presentation.slides.add_slide(presentation.slide_layouts[1])
    slide.shapes.title.text = "经营情况"
    slide.placeholders[1].text = "累计服务客户超过100万户"
    table_shape = slide.shapes.add_table(1, 2, 0, 0, 4_000_000, 1_000_000)
    table_shape.table.cell(0, 0).text = "指标"
    table_shape.table.cell(0, 1).text = "100万户"
    presentation.save(path)


def _make_pptx_with_nested_paragraphs(path: Path) -> None:
    pptx = pytest.importorskip("pptx")
    presentation = pptx.Presentation()
    slide = presentation.slides.add_slide(presentation.slide_layouts[6])
    text_box = slide.shapes.add_textbox(0, 0, 4_000_000, 2_000_000)
    frame = text_box.text_frame
    frame.clear()
    parent = frame.paragraphs[0]
    parent.text = "1、一级"
    child = frame.add_paragraph()
    child.level = 1
    child.text = "1、二级"
    frame.add_paragraph()
    parent_two = frame.add_paragraph()
    parent_two.text = "2、一级"
    presentation.save(path)


def _inject_ooxml_parts(
    path: Path,
    *,
    content_type_entries: str,
    parts: dict[str, bytes],
    content_type_replacements: dict[str, str] | None = None,
    relationship_entries: str = "",
) -> None:
    rebuilt = path.with_suffix(".rebuilt")
    relationship_updated = False
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        rebuilt,
        "w",
        compression=zipfile.ZIP_DEFLATED,
    ) as target:
        for info in source.infolist():
            payload = source.read(info.filename)
            if info.filename == "[Content_Types].xml":
                text = payload.decode("utf-8")
                for old, new in (content_type_replacements or {}).items():
                    text = text.replace(old, new)
                payload = text.replace(
                    "</Types>",
                    f"{content_type_entries}</Types>",
                ).encode("utf-8")
            if (
                relationship_entries
                and info.filename == "ppt/slides/_rels/slide1.xml.rels"
            ):
                text = payload.decode("utf-8")
                payload = text.replace(
                    "</Relationships>",
                    f"{relationship_entries}</Relationships>",
                ).encode("utf-8")
                relationship_updated = True
            target.writestr(info, payload)
        if relationship_entries and not relationship_updated:
            target.writestr(
                "ppt/slides/_rels/slide1.xml.rels",
                (
                    '<Relationships xmlns="http://schemas.openxmlformats.org/'
                    'package/2006/relationships">'
                    f"{relationship_entries}</Relationships>"
                ).encode("utf-8"),
            )
        for name, payload in parts.items():
            target.writestr(name, payload)
    rebuilt.replace(path)


def test_document_service_parses_docx_and_persists_complete_artifact(tmp_path):
    input_dir = tmp_path / "input"
    work_dir = tmp_path / "work"
    input_dir.mkdir()
    path = input_dir / "材料.docx"
    _make_docx(path)

    artifact = DocumentService().parse(
        path,
        allowed_root=input_dir,
        work_dir=work_dir,
    )

    assert artifact.format == DocumentFormat.DOCX
    assert "测试标题" in artifact.full_text
    assert "第一段正文" in artifact.full_text
    assert "指标\t100亿元" in artifact.full_text
    assert any(block.kind == "header" and "页眉文字" in block.text for block in artifact.blocks)
    assert any(block.kind == "table" for block in artifact.blocks)
    assert artifact.sha256

    stored = list((work_dir / "documents").glob("*/document.json"))
    assert len(stored) == 1
    payload = json.loads(stored[0].read_text(encoding="utf-8"))
    assert payload["format"] == "docx"
    assert payload["full_text"] == artifact.full_text
    assert payload["source"]["original_name"] == "材料.docx"


def test_document_service_marks_embedded_docx_images_at_original_position(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "科技部素材.docx"
    document = Document()
    document.add_paragraph("图片前的事实。")
    image_paragraph = document.add_paragraph()
    image_paragraph.add_run().add_picture(_tiny_png(), width=Inches(0.1))
    document.add_paragraph("图片后的事实。")
    document.save(path)

    artifact = DocumentService().parse(path, allowed_root=input_dir, work_dir=tmp_path / "work")

    reminder = "【提醒：科技部素材含图片，请评估是否需要】"
    assert reminder in artifact.full_text
    assert artifact.full_text.index("图片前的事实") < artifact.full_text.index(reminder)
    assert artifact.full_text.index(reminder) < artifact.full_text.index("图片后的事实")
    assert len(artifact.assets) == 1
    assert Path(artifact.assets[0].path).exists()
    assert {warning.code for warning in artifact.warnings} == {"embedded_image_unread"}
    material = artifact.to_material()
    assert material["asset_count"] == 1
    assert material["warning_codes"] == ["embedded_image_unread"]


def test_document_service_marks_textless_pdf_pages_for_ocr(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "扫描材料.pdf"
    _make_blank_pdf(path, pages=2)

    artifact = DocumentService().parse(path, allowed_root=input_dir, work_dir=tmp_path / "work")

    assert artifact.format == DocumentFormat.PDF
    assert artifact.page_count == 2
    assert artifact.full_text == ""
    assert {warning.code for warning in artifact.warnings} == {"ocr_required"}
    assert artifact.warnings[0].locations == ("page:1", "page:2")


def test_document_service_parses_pptx_with_slide_locations_and_tables(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "汇报材料.pptx"
    _make_pptx(path)

    artifact = DocumentService().parse(path, allowed_root=input_dir, work_dir=tmp_path / "work")

    assert artifact.format == DocumentFormat.PPTX
    assert artifact.page_count == 1
    assert "经营情况" in artifact.full_text
    assert "累计服务客户超过100万户" in artifact.full_text
    assert "指标\t100万户" in artifact.full_text
    assert all(block.location.startswith("slide:1") for block in artifact.blocks)
    assert any(block.kind == "table" for block in artifact.blocks)


def test_document_service_preserves_pptx_paragraph_levels_as_indentation(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "分级序号.pptx"
    _make_pptx_with_nested_paragraphs(path)

    artifact = DocumentService().parse(
        path,
        allowed_root=input_dir,
        work_dir=tmp_path / "work",
    )

    text_block = next(block for block in artifact.blocks if block.kind == "text")
    assert text_block.text.splitlines() == ["1、一级", "  1、二级", "", "2、一级"]


def test_document_service_allows_inert_embedded_xlsb_chart_data(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "含图表数据.pptx"
    _make_pptx(path)
    _inject_ooxml_parts(
        path,
        content_type_entries=(
            '<Default Extension="xlsb" '
            'ContentType="application/vnd.ms-excel.sheet.binary.macroEnabled.12"/>'
        ),
        parts={"ppt/embeddings/chart-data.xlsb": b"inert chart cache"},
        relationship_entries=(
            '<Relationship Id="rIdChartData" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/'
            'relationships/package" Target="../embeddings/chart-data.xlsb"/>'
        ),
    )

    artifact = DocumentService().parse(
        path,
        allowed_root=input_dir,
        work_dir=tmp_path / "work",
    )

    assert artifact.format == DocumentFormat.PPTX


def test_document_service_allows_canonical_embedded_xlsb_override(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "含覆盖声明图表数据.pptx"
    _make_pptx(path)
    _inject_ooxml_parts(
        path,
        content_type_entries=(
            '<Override PartName="/ppt/embeddings/chart-data.xlsb" '
            'ContentType="application/vnd.ms-excel.sheet.binary.macroEnabled.12"/>'
        ),
        parts={"ppt/embeddings/chart-data.xlsb": b"inert chart cache"},
        relationship_entries=(
            '<Relationship Id="rIdChartData" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/'
            'relationships/package" Target="../embeddings/chart-data.xlsb"/>'
        ),
    )

    artifact = DocumentService().parse(
        path,
        allowed_root=input_dir,
        work_dir=tmp_path / "work",
    )

    assert artifact.format == DocumentFormat.PPTX


def test_document_service_rejects_xlsb_outside_ppt_embeddings(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "异常位置二进制工作簿.pptx"
    _make_pptx(path)
    _inject_ooxml_parts(
        path,
        content_type_entries=(
            '<Default Extension="xlsb" '
            'ContentType="application/vnd.ms-excel.sheet.binary.macroEnabled.12"/>'
        ),
        parts={"ppt/custom/chart-data.xlsb": b"unexpected workbook"},
    )

    with pytest.raises(DocumentSecurityError, match="包含宏"):
        DocumentService().parse(
            path,
            allowed_root=input_dir,
            work_dir=tmp_path / "work",
        )


@pytest.mark.parametrize(
    "part_name",
    (
        "ppt/embeddings/%2e%2e%2fevil.xlsb",
        "ppt/embeddings/..\\evil.xlsb",
    ),
)
def test_document_service_rejects_noncanonical_xlsb_embedding_paths(
    tmp_path,
    part_name,
):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "非规范嵌入路径.pptx"
    _make_pptx(path)
    _inject_ooxml_parts(
        path,
        content_type_entries=(
            f'<Override PartName="/{part_name}" '
            'ContentType="application/vnd.ms-excel.sheet.binary.macroEnabled.12"/>'
        ),
        parts={part_name: b"unexpected workbook"},
    )

    with pytest.raises(DocumentSecurityError, match="包含宏"):
        DocumentService().parse(
            path,
            allowed_root=input_dir,
            work_dir=tmp_path / "work",
        )


def test_document_service_still_rejects_actual_vba_project_in_pptx(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "含宏项目.pptx"
    _make_pptx(path)
    _inject_ooxml_parts(
        path,
        content_type_entries=(
            '<Override PartName="/ppt/vbaProject.bin" '
            'ContentType="application/vnd.ms-office.vbaProject"/>'
        ),
        parts={"ppt/vbaProject.bin": b"macro project"},
    )

    with pytest.raises(DocumentSecurityError, match="包含宏"):
        DocumentService().parse(
            path,
            allowed_root=input_dir,
            work_dir=tmp_path / "work",
        )


def test_document_service_rejects_macro_enabled_presentation_main_part(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "伪装主文档宏.pptx"
    _make_pptx(path)
    safe_type = (
        "application/vnd.openxmlformats-officedocument.presentationml."
        "presentation.main+xml"
    )
    macro_type = "application/vnd.ms-powerpoint.presentation.macroEnabled.main+xml"
    _inject_ooxml_parts(
        path,
        content_type_entries="",
        parts={},
        content_type_replacements={safe_type: macro_type},
    )

    with pytest.raises(DocumentSecurityError, match="包含宏"):
        DocumentService().parse(
            path,
            allowed_root=input_dir,
            work_dir=tmp_path / "work",
        )


def test_document_service_rejects_noncanonical_main_override_bypass(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "非规范主部件.pptx"
    _make_pptx(path)
    _inject_ooxml_parts(
        path,
        content_type_entries=(
            '<Default Extension="xml" '
            'ContentType="application/vnd.ms-powerpoint.presentation.'
            'macroEnabled.main+xml"/>'
        ),
        parts={},
        content_type_replacements={
            'PartName="/ppt/presentation.xml"':
            'PartName="//ppt/presentation.xml"'
        },
    )

    with pytest.raises(DocumentSecurityError):
        DocumentService().parse(
            path,
            allowed_root=input_dir,
            work_dir=tmp_path / "work",
        )


def test_document_service_rejects_fake_content_type_element_bypass(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "伪元素主部件.pptx"
    _make_pptx(path)
    safe_type = (
        "application/vnd.openxmlformats-officedocument.presentationml."
        "presentation.main+xml"
    )
    _inject_ooxml_parts(
        path,
        content_type_entries=(
            '<Default Extension="xml" '
            'ContentType="application/vnd.ms-powerpoint.presentation.'
            'macroEnabled.main+xml"/>'
            f'<Fake PartName="/ppt/presentation.xml" ContentType="{safe_type}"/>'
        ),
        parts={},
        content_type_replacements={
            'PartName="/ppt/presentation.xml"':
            'PartName="/ppt/not-the-main-presentation.xml"'
        },
    )

    with pytest.raises(DocumentSecurityError):
        DocumentService().parse(
            path,
            allowed_root=input_dir,
            work_dir=tmp_path / "work",
        )


def test_document_service_rejects_duplicate_main_overrides(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "重复主部件声明.pptx"
    _make_pptx(path)
    safe_type = (
        "application/vnd.openxmlformats-officedocument.presentationml."
        "presentation.main+xml"
    )
    _inject_ooxml_parts(
        path,
        content_type_entries=(
            '<Override PartName="/ppt/presentation.xml" '
            f'ContentType="{safe_type}"/>'
        ),
        parts={},
    )

    with pytest.raises(DocumentSecurityError):
        DocumentService().parse(
            path,
            allowed_root=input_dir,
            work_dir=tmp_path / "work",
        )


def test_document_service_rejects_other_macro_enabled_embedded_types(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "嵌入宏工作簿.pptx"
    _make_pptx(path)
    _inject_ooxml_parts(
        path,
        content_type_entries=(
            '<Default Extension="xlsm" '
            'ContentType="application/vnd.ms-excel.sheet.macroEnabled.12"/>'
        ),
        parts={"ppt/embeddings/embedded.xlsm": b"macro-enabled workbook"},
    )

    with pytest.raises(DocumentSecurityError, match="包含宏"):
        DocumentService().parse(
            path,
            allowed_root=input_dir,
            work_dir=tmp_path / "work",
        )


def test_document_service_rejects_vba_project_relationship(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "宏关系.pptx"
    _make_pptx(path)
    _inject_ooxml_parts(
        path,
        content_type_entries="",
        parts={},
        relationship_entries=(
            '<Relationship Id="rIdVba" '
            'Type="http://schemas.microsoft.com/office/2006/relationships/'
            'vbaProject" Target="../hidden.bin"/>'
        ),
    )

    with pytest.raises(DocumentSecurityError, match="包含宏"):
        DocumentService().parse(
            path,
            allowed_root=input_dir,
            work_dir=tmp_path / "work",
        )


def test_document_service_does_not_allow_embedded_xlsb_in_docx(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "嵌入二进制工作簿.docx"
    _make_docx(path)
    _inject_ooxml_parts(
        path,
        content_type_entries=(
            '<Default Extension="xlsb" '
            'ContentType="application/vnd.ms-excel.sheet.binary.macroEnabled.12"/>'
        ),
        parts={"word/embeddings/chart-data.xlsb": b"embedded workbook"},
    )

    with pytest.raises(DocumentSecurityError, match="包含宏"):
        DocumentService().parse(
            path,
            allowed_root=input_dir,
            work_dir=tmp_path / "work",
        )


def test_pptx_chart_parser_reports_unreadable_series_values():
    class BrokenSeries:
        name = "客户数"

        @property
        def values(self):
            raise ValueError("unreadable chart cache")

    class BrokenChart:
        has_title = False
        series = (BrokenSeries(),)

    text, warnings = _chart_text(
        BrokenChart(),
        slide_number=2,
        location="slide:2/shape:3",
    )

    assert text == "客户数"
    assert len(warnings) == 1
    assert warnings[0].code == "pptx_chart_values_unreadable"
    assert warnings[0].message == "第2页图表“客户数”数据未完整读取"
    assert warnings[0].locations == ("slide:2/shape:3",)


def test_document_service_rejects_extension_signature_mismatch(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    fake_pdf = input_dir / "伪造.pdf"
    fake_pdf.write_bytes(b"not-a-pdf")

    with pytest.raises(DocumentSecurityError, match="文件内容与 PDF 格式不一致"):
        DocumentService().parse(fake_pdf, allowed_root=input_dir, work_dir=tmp_path / "work")


def test_document_service_rejects_paths_outside_current_task(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    outside = tmp_path / "outside.pdf"
    _make_blank_pdf(outside, pages=1)

    with pytest.raises(DocumentSecurityError, match="当前任务目录之外"):
        DocumentService().parse(outside, allowed_root=input_dir, work_dir=tmp_path / "work")


def test_document_service_rejects_html_until_scope_is_reopened(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    html = input_dir / "材料.html"
    html.write_text("<html><body>正文</body></html>", encoding="utf-8")

    with pytest.raises(DocumentSecurityError, match="暂不支持"):
        DocumentService().parse(html, allowed_root=input_dir, work_dir=tmp_path / "work")


def test_document_service_rejects_file_over_configured_limit(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "超限.pdf"
    path.write_bytes(b"%PDF-" + b"0" * 32)

    with pytest.raises(DocumentSecurityError, match="大小上限"):
        DocumentService(max_file_bytes=16).parse(
            path,
            allowed_root=input_dir,
            work_dir=tmp_path / "work",
        )


def test_document_service_rejects_intermediate_output_outside_task_work(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "材料.docx"
    _make_docx(path)

    with pytest.raises(ValueError, match="work 目录"):
        DocumentService().parse(
            path,
            allowed_root=input_dir,
            work_dir=tmp_path / "output",
        )


def test_document_service_rejects_suspicious_ooxml_compression_ratio(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "异常.docx"
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Override PartName="/word/document.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>",
        )
        archive.writestr("word/document.xml", "A" * 2_000_000)

    with pytest.raises(DocumentSecurityError, match="压缩比异常"):
        DocumentService(max_compression_ratio=100).parse(
            path,
            allowed_root=input_dir,
            work_dir=tmp_path / "work",
        )


def test_prompt_material_samples_across_long_document_instead_of_prefix_only(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "长材料.docx"
    document = Document()
    for index in range(1, 31):
        document.add_paragraph(f"第{index}段-" + str(index) * 80)
    document.save(path)

    artifact = DocumentService().parse(path, allowed_root=input_dir, work_dir=tmp_path / "work")
    material = artifact.to_material(max_chars=1200)

    assert material["content_complete"] is False
    assert "第1段" in material["text"]
    assert "第30段" in material["text"]
    assert "中间内容已按位置抽样" in material["text"]
    assert Path(str(material["artifact_path"])).is_relative_to(tmp_path / "work")


def test_long_docx_material_sampling_never_drops_image_reminder(tmp_path):
    input_dir = tmp_path / "input"
    input_dir.mkdir()
    path = input_dir / "科技部素材.docx"
    document = Document()
    for index in range(1, 41):
        document.add_paragraph(f"第{index}段-" + str(index) * 80)
        if index == 13:
            image_paragraph = document.add_paragraph()
            image_paragraph.add_run().add_picture(_tiny_png(), width=Inches(0.1))
    document.save(path)

    artifact = DocumentService().parse(path, allowed_root=input_dir, work_dir=tmp_path / "work")
    material = artifact.to_material(max_chars=1200)

    assert "【提醒：科技部素材含图片，请评估是否需要】" in material["text"]
