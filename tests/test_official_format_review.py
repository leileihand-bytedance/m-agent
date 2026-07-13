"""独立公文格式审核规则测试。"""

from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.review.document_type import DocumentType, detect_document_type
from app.review.error_marker import mark_errors_in_docx
from app.review.official_format_checker import review_official_format


def _set_run_format(run, *, font: str, size: float, bold: bool) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold


def _add_formatted_paragraph(
    document: Document,
    text: str,
    *,
    font: str,
    size: float = 16,
    bold: bool = False,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_pt: float = 32,
):
    # 故意全部使用 Normal，验证审核只看实际格式，不依赖 Word 样式名称。
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = alignment
    paragraph.paragraph_format.first_line_indent = Pt(first_line_pt)
    paragraph.paragraph_format.line_spacing = Pt(29)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    _set_run_format(run, font=font, size=size, bold=bold)
    return paragraph


def _build_template_like_document(path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.175)
    section.right_margin = Cm(3.175)

    _add_formatted_paragraph(
        document,
        "相关情况报告",
        font="宋体",
        size=18,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_pt=0,
    )
    _add_formatted_paragraph(document, "这是公文正文第一段。", font="仿宋")
    _add_formatted_paragraph(document, "一、总体情况", font="黑体")
    _add_formatted_paragraph(document, "这是一级标题后的正文。", font="仿宋")
    _add_formatted_paragraph(document, "（一）经营情况", font="楷体", bold=True)
    _add_formatted_paragraph(document, "这是二级标题后的正文。", font="仿宋")
    _add_formatted_paragraph(document, "1.重点工作", font="仿宋", bold=True)
    _add_formatted_paragraph(document, "这是三级标题后的正文。", font="仿宋")
    document.save(path)


def test_official_format_accepts_actual_format_without_heading_styles(tmp_path: Path):
    source = tmp_path / "符合模板.docx"
    _build_template_like_document(source)

    result = review_official_format(source, source.name)

    assert result.findings == []
    assert result.passed_rules == result.total_rules


def test_official_format_is_never_selected_from_filename_or_content():
    doc_type = detect_document_type(
        "公文格式审核.docx",
        ["关于有关情况的报告", "一、总体情况"],
    )

    assert doc_type == DocumentType.GENERAL


def test_official_format_reports_actual_font_size_spacing_and_page_errors(tmp_path: Path):
    source = tmp_path / "格式错误.docx"
    _build_template_like_document(source)
    document = Document(source)
    document.sections[0].left_margin = Cm(2.0)

    title = document.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_format(title.runs[0], font="黑体", size=16, bold=False)

    heading1 = document.paragraphs[2]
    _set_run_format(heading1.runs[0], font="仿宋", size=16, bold=False)

    body = document.paragraphs[3]
    body.paragraph_format.line_spacing = Pt(20)

    heading2 = document.paragraphs[4]
    _set_run_format(heading2.runs[0], font="楷体", size=14, bold=False)
    document.save(source)

    result = review_official_format(source, source.name)
    rule_ids = {finding.rule_id for finding in result.findings}

    assert "official-format-page" in rule_ids
    assert "official-format-title" in rule_ids
    assert "official-format-heading1" in rule_ids
    assert "official-format-heading2" in rule_ids
    assert "official-format-body" in rule_ids
    assert any("29pt" in finding.description for finding in result.findings)


def test_official_format_does_not_treat_normal_body_as_heading(tmp_path: Path):
    source = tmp_path / "正文含数字.docx"
    _build_template_like_document(source)
    document = Document(source)
    paragraph = document.paragraphs[-1]
    paragraph.text = "2026年计划完成3项重点工作。"
    _set_run_format(paragraph.runs[0], font="仿宋", size=16, bold=False)
    document.save(source)

    result = review_official_format(source, source.name)

    assert result.findings == []


def test_official_format_findings_can_be_marked_in_returned_docx(tmp_path: Path):
    source = tmp_path / "标题格式错误.docx"
    marked = tmp_path / "marked_标题格式错误.docx"
    _build_template_like_document(source)
    document = Document(source)
    _set_run_format(document.paragraphs[0].runs[0], font="黑体", size=16, bold=False)
    document.save(source)

    result = review_official_format(source, source.name)
    mark_errors_in_docx(source, marked, result.findings)

    assert marked.exists()
    with zipfile.ZipFile(marked) as archive:
        comments = archive.read("word/comments.xml").decode("utf-8")
    assert "标题" in comments
    assert "宋体" in comments
