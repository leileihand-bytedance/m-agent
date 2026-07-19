from datetime import date
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from skills.shenyinxie_news import docx_output
from skills.shenyinxie_news.docx_output import write_shenyinxie_docx
from skills.shenyinxie_news.schema import SelectedArticle


def _article(**kwargs) -> SelectedArticle:
    defaults = {
        "title": "微众银行发布年报",
        "media_name": "人民网",
        "publish_date": "2026-07-15",
        "body": "2026年7月15日，微众银行发布年报，营收增长20%。" * 5,
        "original_url": "https://people.com.cn/1",
    }
    defaults.update(kwargs)
    return SelectedArticle(**defaults)


def test_docx_output_without_template_uses_scratch_format(tmp_path):
    output_dir = tmp_path / "out"
    articles = [_article()]

    path = write_shenyinxie_docx(
        title="深圳银行业协会工作动态（2026年7月第2期）",
        period_start=date(2026, 7, 16),
        period_end=date(2026, 7, 29),
        issue_number="2026-14",
        articles=articles,
        output_dir=output_dir,
        template_path=tmp_path / "nonexistent.docx",
    )

    assert path.exists()
    assert path.name == "【深银协】微众银行2026年7月第2期信息动态.docx"

    doc = Document(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "深圳银行业协会工作动态" in full_text
    assert "动态一" in full_text
    assert "2026年7月16日-7月29日" in full_text
    assert "微众银行发布年报" in full_text
    assert "https://people.com.cn/1" in full_text


def test_docx_output_with_placeholder_template_replaces_articles(tmp_path):
    template = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("{{TITLE}}")
    doc.add_paragraph("{{PERIOD_RANGE}}")
    doc.add_paragraph("{{ARTICLE_1}}")
    doc.add_paragraph("{{ARTICLE_2}}")
    doc.add_paragraph("{{ARTICLE_3}}")
    doc.save(str(template))

    articles = [
        _article(title="第一篇", body="正文一" * 5),
        _article(title="第二篇", body="正文二" * 5),
    ]

    path = write_shenyinxie_docx(
        title="T",
        period_start=date(2026, 7, 16),
        period_end=date(2026, 7, 29),
        issue_number="2026-14",
        articles=articles,
        output_dir=tmp_path / "out",
        template_path=template,
    )

    assert path.exists()
    doc = Document(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "T" in full_text
    assert "2026年7月16日-7月29日" in full_text
    assert "第一篇" in full_text
    assert "第二篇" in full_text
    assert "正文一" in full_text
    assert "正文二" in full_text
    # 第三篇占位符应被替换为空
    assert "{{ARTICLE_3}}" not in full_text


def test_default_template_keeps_reference_page_and_role_styles():
    template = (
        Path(__file__).resolve().parents[1]
        / "skills"
        / "shenyinxie_news"
        / "assets"
        / "shenyinxie-template.docx"
    )
    doc = Document(str(template))
    section = doc.sections[0]
    texts = [paragraph.text for paragraph in doc.paragraphs]

    assert section.page_width.cm == pytest.approx(21.0, abs=0.02)
    assert section.page_height.cm == pytest.approx(29.7, abs=0.02)
    assert section.left_margin.cm == pytest.approx(3.17, abs=0.02)
    assert section.right_margin.cm == pytest.approx(3.17, abs=0.02)
    assert "{{TITLE}}" in texts
    assert "{{PERIOD_RANGE}}" in texts
    assert "{{ARTICLE_1_TITLE}}" in texts
    assert "{{ARTICLE_1_BODY}}" in texts
    assert "{{ARTICLE_1_SOURCE}}" in texts
    assert "金融“春雨”润鹏城" not in "\n".join(texts)

    title = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "{{TITLE}}")
    assert title.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert title.runs[0].font.size.pt == pytest.approx(16.0)
    assert title.runs[0].bold is True

    heading = next(
        paragraph for paragraph in doc.paragraphs if paragraph.text == "{{ARTICLE_1_TITLE}}"
    )
    assert heading.style.name == "Heading 2"
    assert heading.runs[0].font.name == "黑体"
    assert heading.runs[0].font.size.pt == pytest.approx(14.0)

    body = next(
        paragraph for paragraph in doc.paragraphs if paragraph.text == "{{ARTICLE_1_BODY}}"
    )
    assert body.style.name == "Normal (Web)"
    assert body.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert body.paragraph_format.first_line_indent.pt == pytest.approx(24.0, abs=0.2)
    assert body.paragraph_format.line_spacing == pytest.approx(1.5)


def test_default_template_output_uses_real_paragraphs_and_preserves_styles(tmp_path):
    article = _article(
        title="微众银行科技创新取得新成果",
        body="第一段正文，介绍微众银行科技创新成果。\n\n第二段正文，介绍数字普惠金融成效。",
    )

    path = write_shenyinxie_docx(
        title="微众银行2026年7月第1期信息动态",
        period_start=date(2026, 7, 1),
        period_end=date(2026, 7, 15),
        issue_number="2026-13",
        articles=[article],
        output_dir=tmp_path / "out",
    )

    doc = Document(str(path))
    texts = [paragraph.text for paragraph in doc.paragraphs]
    assert "微众银行2026年7月第1期信息动态" in texts
    assert "（2026年7月1日-7月15日）" in texts
    assert "【动态一】微众银行科技创新取得新成果" in texts
    assert "第一段正文，介绍微众银行科技创新成果。" in texts
    assert "第二段正文，介绍数字普惠金融成效。" in texts
    assert not any("\n" in text for text in texts)
    assert not any("{{" in text for text in texts)

    title = next(
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text == "微众银行2026年7月第1期信息动态"
    )
    assert title.runs[0].font.size.pt == pytest.approx(16.0)
    assert title.runs[0].bold is True

    heading = next(
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text == "【动态一】微众银行科技创新取得新成果"
    )
    assert heading.style.name == "Heading 2"
    assert heading.runs[0].font.name == "Hiragino Sans GB"

    body = next(
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text == "第一段正文，介绍微众银行科技创新成果。"
    )
    assert body.style.name == "Normal (Web)"
    assert body.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert body.paragraph_format.first_line_indent.pt == pytest.approx(24.0, abs=0.2)

    source = next(paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("来源：人民网"))
    assert "发布时间：2026年7月15日" in source.text
    assert "原文链接：https://people.com.cn/1" in source.text
    assert source.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert source.paragraph_format.first_line_indent.pt == pytest.approx(0.0, abs=0.2)
    assert any(
        relationship.target_ref == "https://people.com.cn/1"
        and relationship.is_external
        for relationship in doc.part.rels.values()
    )


def test_default_template_output_maps_chinese_fonts_for_rendering(tmp_path):
    path = write_shenyinxie_docx(
        title="微众银行2026年7月第1期信息动态",
        period_start=date(2026, 7, 1),
        period_end=date(2026, 7, 15),
        issue_number="2026-13",
        articles=[_article()],
        output_dir=tmp_path / "out",
    )

    doc = Document(str(path))
    heading = next(paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("【动态一】"))
    heading_fonts = heading.runs[0]._element.get_or_add_rPr().get_or_add_rFonts()
    body_style_fonts = doc.styles["Normal (Web)"]._element.get_or_add_rPr().get_or_add_rFonts()

    assert heading_fonts.get(qn("w:eastAsia")) == "Hiragino Sans GB"
    assert body_style_fonts.get(qn("w:eastAsia")) == "Songti SC"


def test_docx_output_removes_markdown_front_matter_defensively(tmp_path):
    article = _article(
        body=(
            "---\n"
            "title: 微众银行AI算力增长3.5倍\n"
            "source: 上海证券报\n"
            "canonical_url: https://example.com/article\n"
            "---\n"
            "微众银行持续推进AI原生银行建设。"
        )
    )

    path = write_shenyinxie_docx(
        title="微众银行2026年7月第1期信息动态",
        period_start=date(2026, 7, 1),
        period_end=date(2026, 7, 15),
        issue_number="2026-13",
        articles=[article],
        output_dir=tmp_path / "out",
    )

    full_text = "\n".join(paragraph.text for paragraph in Document(str(path)).paragraphs)
    assert "title:" not in full_text
    assert "source:" not in full_text
    assert "canonical_url:" not in full_text
    assert "微众银行持续推进AI原生银行建设。" in full_text


def test_docx_output_with_blank_template_falls_back_to_scratch(tmp_path):
    """占位模板不含约定占位符时，应视为空白占位并按代码新建文档。"""
    template = tmp_path / "blank.docx"
    Document().save(str(template))

    articles = [_article()]
    path = write_shenyinxie_docx(
        title="深圳银行业协会工作动态（2026年7月第2期）",
        period_start=date(2026, 7, 16),
        period_end=date(2026, 7, 29),
        issue_number="2026-14",
        articles=articles,
        output_dir=tmp_path / "out",
        template_path=template,
    )

    doc = Document(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "深圳银行业协会工作动态" in full_text
    assert "动态一" in full_text


def test_default_template_missing_stops_instead_of_rebuilding_layout(tmp_path, monkeypatch):
    missing_template = tmp_path / "missing-template.docx"
    monkeypatch.setattr(docx_output, "_default_template_path", lambda: missing_template)

    with pytest.raises(RuntimeError, match="深银协正式模板不可用"):
        write_shenyinxie_docx(
            title="微众银行2026年7月第1期信息动态",
            period_start=date(2026, 7, 1),
            period_end=date(2026, 7, 15),
            issue_number="2026-13",
            articles=[_article()],
            output_dir=tmp_path / "out",
        )


def test_docx_output_cross_month_period_range(tmp_path):
    articles = [_article()]
    path = write_shenyinxie_docx(
        title="T",
        period_start=date(2026, 7, 16),
        period_end=date(2026, 8, 15),
        issue_number="2026-15",
        articles=articles,
        output_dir=tmp_path / "out",
        template_path=tmp_path / "nonexistent.docx",
    )

    doc = Document(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "2026年7月16日-8月15日" in full_text


def test_docx_output_includes_all_three_articles(tmp_path):
    articles = [
        _article(title="第一篇"),
        _article(title="第二篇"),
        _article(title="第三篇"),
    ]
    path = write_shenyinxie_docx(
        title="T",
        period_start=date(2026, 7, 16),
        period_end=date(2026, 7, 29),
        issue_number="2026-14",
        articles=articles,
        output_dir=tmp_path / "out",
        template_path=tmp_path / "nonexistent.docx",
    )

    doc = Document(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "动态一" in full_text
    assert "动态二" in full_text
    assert "动态三" in full_text
    assert "第一篇" in full_text
    assert "第二篇" in full_text
    assert "第三篇" in full_text


def test_docx_output_path_issue_number_formatting(tmp_path):
    path = write_shenyinxie_docx(
        title="T",
        period_start=date(2026, 7, 16),
        period_end=date(2026, 7, 29),
        issue_number="2026-14",
        articles=[_article()],
        output_dir=tmp_path / "out",
        template_path=tmp_path / "nonexistent.docx",
    )
    assert path.name == "【深银协】微众银行2026年7月第2期信息动态.docx"


def test_docx_output_removes_web_page_chrome_noise(tmp_path):
    article = _article(
        title="科技创新助推数字化金融普惠发展",
        body=(
            "2026-07-11 07:00\n+\n-\n第07版：特别报道\n本版新闻\n"
            "人民日报 2026年07月11日\nSat\n科技创新助推数字化金融普惠发展\n"
            "微众银行党委书记 李南青 《人民日报》（ 2026年07月11日 第\u00a007\u00a0版）\n"
            "微众银行通过科技创新推动数字普惠金融发展。"
        )
    )

    path = write_shenyinxie_docx(
        title="微众银行2026年7月第1期信息动态",
        period_start=date(2026, 7, 1),
        period_end=date(2026, 7, 15),
        issue_number="2026-13",
        articles=[article],
        output_dir=tmp_path / "out",
    )

    texts = [paragraph.text for paragraph in Document(str(path)).paragraphs]
    assert "微众银行通过科技创新推动数字普惠金融发展。" in texts
    assert "微众银行党委书记 李南青" in texts
    assert "科技创新助推数字化金融普惠发展" not in texts
    assert "+" not in texts
    assert "-" not in texts
    assert "Sat" not in texts
    assert not any(text.startswith("第07版") for text in texts)
    assert not any(text.startswith("人民日报 2026年") for text in texts)
    assert not any("第\u00a007\u00a0版" in text for text in texts)


def test_docx_excerpt_includes_source_title_and_editor_note(tmp_path):
    article = _article(
        title="微众银行连续两年实施利润分配",
        source_title="民营银行利润分配观察",
        content_mode="extract",
        editor_note="说明：本文根据原报道中微众银行相关内容摘编。",
    )

    path = write_shenyinxie_docx(
        title="T",
        period_start=date(2026, 7, 1),
        period_end=date(2026, 7, 15),
        issue_number="2026-14",
        articles=[article],
        output_dir=tmp_path / "out",
        template_path=tmp_path / "nonexistent.docx",
    )

    doc = Document(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "原报道标题：民营银行利润分配观察" in full_text
    assert "原文链接：https://people.com.cn/1" in full_text
    assert "说明：本文根据原报道中微众银行相关内容摘编。" in full_text


def test_docx_full_text_does_not_include_excerpt_disclosure(tmp_path):
    path = write_shenyinxie_docx(
        title="T",
        period_start=date(2026, 7, 1),
        period_end=date(2026, 7, 15),
        issue_number="2026-14",
        articles=[_article(content_mode="full_text", source_title="微众银行发布年报")],
        output_dir=tmp_path / "out",
        template_path=tmp_path / "nonexistent.docx",
    )

    doc = Document(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "原报道标题：" not in full_text
    assert "摘编" not in full_text


def test_docx_placeholder_template_includes_excerpt_disclosure(tmp_path):
    template = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("{{TITLE}}")
    doc.add_paragraph("{{ARTICLE_1}}")
    doc.save(str(template))
    article = _article(
        title="微众银行连续两年实施利润分配",
        source_title="民营银行利润分配观察",
        content_mode="extract",
        editor_note="说明：本文根据原报道中微众银行相关内容摘编。",
    )

    path = write_shenyinxie_docx(
        title="T",
        period_start=date(2026, 7, 1),
        period_end=date(2026, 7, 15),
        issue_number="2026-14",
        articles=[article],
        output_dir=tmp_path / "out",
        template_path=template,
    )

    full_text = "\n".join(p.text for p in Document(str(path)).paragraphs)
    assert "原报道标题：民营银行利润分配观察" in full_text
    assert "原文链接：https://people.com.cn/1" in full_text
    assert "说明：本文根据原报道中微众银行相关内容摘编。" in full_text
